"""CERA Richiesta di Parere — ARISE feasibility study at REHELAB Savona.

Style mirrors the prior Marco Testa / DINOGMI CERA submissions: concise plain
Italian, short paragraphs, healthy-volunteer observational data collection,
no medical-device language, single UNIGE site (REHELAB, Campus di Savona).
Innovina S.r.l. retained as technical-industrial partner (software supplier),
not as a clinical site.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2")
OUT  = ROOT / "ARISE_CERA_Richiesta_Parere.docx"

CHECKED, UNCHECKED = "☒", "☐"

def heading0(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16)

def heading1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13)

def label(doc, text, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(11)

def body(doc, text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(11)

def field(doc, key, val):
    p = doc.add_paragraph()
    a = p.add_run(f"{key}: "); a.bold = True; a.font.size = Pt(11)
    b = p.add_run(val); b.font.size = Pt(11)

def checks(doc, items):
    for c, t in items:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f"{CHECKED if c else UNCHECKED}  {t}"); r.font.size = Pt(11)

def spacer(doc, n=1):
    for _ in range(n): doc.add_paragraph()

doc = Document()
s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)

# Header
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UNIVERSITÀ DEGLI STUDI DI GENOVA"); r.bold = True; r.font.size = Pt(14)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DIPARTIMENTO DI NEUROSCIENZE, RIABILITAZIONE, OFTALMOLOGIA, "
              "GENETICA E SCIENZE MATERNO-INFANTILI (DINOGMI)"); r.font.size = Pt(11)
spacer(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Al Presidente\nComitato Etico per la Ricerca di Ateneo\n"
              "Università degli Studi di Genova"); r.font.size = Pt(11)
spacer(doc)
heading0(doc, "RICHIESTA DI PARERE"); spacer(doc)

# 1. Dati generali
heading1(doc, "1. DATI GENERALI RELATIVI AL PROGETTO DI RICERCA")
field(doc, "Titolo del progetto",
      "\"ARISE: A feasibility study on markerless video-based biomechanical "
      "assessment of the Sit-to-Stand task in healthy adults and older adults\".")
field(doc, "Acronimo", "ARISE")
field(doc, "Responsabile del progetto",
      "Prof. Marco Testa, Professore Associato presso il Dipartimento di "
      "Neuroscienze, Riabilitazione, Oftalmologia, Genetica e Scienze "
      "Materno-Infantili (DINOGMI) dell'Università degli Studi di Genova.")
field(doc, "Indirizzo e-mail", "marco.testa@unige.it")
field(doc, "Recapito telefonico", "+39 331 2611548")
spacer(doc)

label(doc, "Componenti del gruppo di ricerca:")
members = [
    ("Testa Marco",       "Professore Associato, Responsabile del progetto (UNIGE)"),
    ("Pavani Paolo",      "Project Manager, Innovina S.r.l."),
    ("Sahraoui Hadil",    "Ingegnere, Innovina S.r.l."),
]
t = doc.add_table(rows=1+len(members), cols=2); t.style = "Light Grid Accent 1"
t.rows[0].cells[0].text = "Cognome e nome"
t.rows[0].cells[1].text = "Ruolo / Affiliazione"
for i, (n_, r_) in enumerate(members, 1):
    t.rows[i].cells[0].text = n_; t.rows[i].cells[1].text = r_
body(doc, "*Innovina S.r.l. partecipa come partner tecnico-industriale del "
          "progetto, fornendo il software di acquisizione e analisi video. "
          "È inquadrata come Responsabile esterno del trattamento ai sensi "
          "dell'Art. 28 GDPR, vincolata da apposito Data Processing "
          "Agreement (DPA) con il Titolare attualmente in corso di "
          "formalizzazione. Il personale di Innovina (Pavani, Sahraoui) "
          "non avrà accesso ai dati identificativi dei partecipanti; "
          "riceverà esclusivamente i dati cinematici delle acquisizioni "
          "in forma pseudonimizzata (etichettati con codice identificativo).",
     italic=True)

spacer(doc)
field(doc, "Sede/i del progetto",
      "Dipartimento di Neuroscienze, Riabilitazione, Oftalmologia, "
      "Genetica e Scienze Materno-Infantili dell'Università degli "
      "Studi di Genova.")
spacer(doc)

checks(doc, [
    (False, "È necessaria l'autorizzazione di altri enti per l'accesso "
            "ai dati o il coinvolgimento dei partecipanti.   NO"),
    (False, "È necessario il consenso di un rappresentante legale (lo "
            "studio non coinvolge minori né soggetti incapaci).   NO"),
])

# 2. Informazioni progetto
spacer(doc)
heading1(doc, "2. INFORMAZIONI RELATIVE AL PROGETTO DI RICERCA")

label(doc, "Tipologia di ricerca")
checks(doc, [
    (True,  "attività pratiche"),
    (True,  "raccolta dati"),
    (False, "analisi retrospettiva"),
    (False, "repertori"),
    (False, "altro"),
])

spacer(doc)
label(doc, "Sinossi del progetto")

label(doc, "Breve stato dell'arte:")
body(doc,
    "Il Sit-to-Stand (STS) è il passaggio dalla posizione seduta a quella "
    "eretta, ed è uno dei principali indicatori della capacità funzionale "
    "degli arti inferiori e del controllo posturale in popolazioni sane "
    "e cliniche. Le valutazioni correnti (Timed Up and Go, Five Times "
    "Sit-to-Stand, 30-second Chair Stand Test [1,2]) si basano sul tempo "
    "totale di esecuzione e su osservazioni qualitative, senza fornire "
    "informazioni quantitative sulla biomeccanica del movimento. Le "
    "tecnologie di laboratorio (pedane di forza, sistemi motion capture) "
    "forniscono misure accurate ma non sono trasferibili al di fuori del "
    "laboratorio. L'analisi del movimento basata su una o più telecamere "
    "RGB senza marker, abilitata dai recenti progressi nella pose "
    "estimation [3], rappresenta un approccio promettente per estendere "
    "la quantificazione biomeccanica del STS, ma richiede una base di "
    "dati di riferimento ottenuti su popolazioni sane per essere "
    "validata e interpretata in maniera affidabile.")

label(doc, "Obiettivi:")
body(doc,
    "1) Raccolta di dati biomeccanici sul Sit-to-Stand in una "
    "popolazione di adulti (>18 anni), utilizzando registrazioni "
    "video da telecamere RGB e/o RGB-D analizzate attraverso "
    "metodologie non invasive di pose estimation senza marker.")
body(doc,
    "2) Caratterizzare i principali indicatori cinematici (KPI) del "
    "movimento: durata di fase, velocità del bacino, angoli "
    "articolari di tronco e ginocchio, simmetria laterale.")
body(doc,
    "3) Validazione dei dati ottenuti attraverso confronto con misure "
    "di riferimento ottenute attraverso sistemi di analisi cinematica "
    "e cinetica del movimento gold-standard validati con metodologie "
    "non invasive.")
body(doc,
    "4) Esplorare la variabilità inter e intra-soggetto del movimento "
    "nelle differenti fasce d'età.")

label(doc, "Metodologia:")
body(doc,
    "Il presente studio è uno studio quasi-sperimentale di fattibilità, "
    "ovvero uno studio osservazionale non controllato né randomizzato "
    "ad un braccio singolo [4].")
body(doc,
    "I partecipanti saranno adulti (>18 anni), di qualsiasi genere, "
    "in grado di eseguire autonomamente l'esercizio Sit-to-Stand "
    "senza assistenza fisica esterna. Saranno esclusi soggetti con "
    "controindicazioni mediche all'esecuzione ripetuta del "
    "Sit-to-Stand, lesioni acute o interventi chirurgici recenti agli "
    "arti inferiori o al tronco (entro 6 settimane), patologie "
    "neurologiche o muscolo-scheletriche che compromettano "
    "l'esecuzione autonoma del movimento, e soggetti con "
    "compromissione cognitiva tale da impedire la comprensione delle "
    "istruzioni. La loro partecipazione sarà volontaria.")
body(doc,
    "Ogni partecipante prenderà parte a 2 sessioni di acquisizione "
    "dati presso i Laboratori di ricerca del Dipartimento DINOGMI "
    "(Campus Universitario di Savona), della durata di 10-15 minuti "
    "ciascuna. Le sessioni sono strutturate come segue:")
body(doc,
    "a) Accoglienza del partecipante, lettura del modulo informativo "
    "e firma del consenso informato.")
body(doc,
    "b) Spiegazione del task da eseguire: posizione seduta su sedia "
    "standard senza braccioli, braccia incrociate al petto, schiena "
    "con angolazione neutrale rispetto alla verticale, anche flesse "
    "a 90°, piedi appoggiati a terra alla larghezza delle anche.")
body(doc,
    "c) Fase di calibrazione iniziale (durata circa 5 secondi): il "
    "partecipante esegue una sequenza di seduta e alzata controllata "
    "per consentire al sistema di calibrarsi sulle caratteristiche "
    "antropometriche individuali (altezza, lunghezze segmentali, "
    "composizione corporea).")
body(doc,
    "d) Esecuzione del protocollo 5xSTS (Five-Times Sit-to-Stand): "
    "cinque ripetizioni consecutive a tempo del passaggio da seduto "
    "a in piedi e ritorno.")
body(doc,
    "e) Pausa di 1 minuto.")
body(doc,
    "f) Esecuzione del protocollo 30-second Chair Stand Test: "
    "esecuzione del maggior numero possibile di ripetizioni del "
    "Sit-to-Stand in 30 secondi.")
body(doc,
    "I dati raccolti saranno utilizzati esclusivamente a fini di "
    "analisi biomeccanica e di ricerca scientifica.")

label(doc, "Analisi dei dati:")
body(doc,
    "L'analisi dei dati seguirà la pipeline sviluppata nell'ambito del "
    "progetto ARISE, articolata in sette stadi.")
body(doc,
    "(1) Anonimizzazione (pseudonimizzazione) alla sorgente. Al momento "
    "della creazione di ogni nuova sessione, il software di acquisizione "
    "fornito da Innovina S.r.l. genera automaticamente un codice "
    "identificativo randomico (es. sbj4112) che etichetta tutti i dati "
    "video e cinematici relativi al/la partecipante. Nessun dato "
    "identificativo diretto (nome, cognome, recapito) viene inserito "
    "nel software. La mappatura codice ↔ identità è conservata "
    "esclusivamente dal Responsabile del progetto (Prof. Marco Testa) "
    "in archivio cifrato presso i Laboratori del Dipartimento DINOGMI, "
    "ai soli fini di esercizio dei diritti GDPR del/della partecipante.")
body(doc,
    "(2) Pre-processing dei video. Ogni registrazione video sarà "
    "processata fotogramma per fotogramma da un sistema di pose "
    "estimation senza marker (MediaPipe Pose Landmarker) per estrarre "
    "33 keypoint anatomici tridimensionali. La serie temporale dei "
    "keypoint sarà sottoposta a filtraggio passa-basso (Butterworth, "
    "frequenza di taglio determinata sulla base della densità spettrale "
    "del segnale) e a rimozione degli outlier basata sul punteggio di "
    "visibilità per fotogramma. La sincronizzazione temporale con la "
    "strumentazione di riferimento (sistemi motion capture, pedane di "
    "forza) avverrà tramite trigger condiviso o cross-correlazione del "
    "segnale di velocità verticale del bacino.")
body(doc,
    "(3) Segmentazione delle ripetizioni e delle fasi. Da ciascuna "
    "sessione saranno isolate le singole ripetizioni del Sit-to-Stand "
    "mediante un algoritmo di peak detection sulla coordinata verticale "
    "del bacino. Ogni ripetizione sarà ulteriormente segmentata nelle "
    "fasi caratteristiche del movimento (seduto, flessione anteriore "
    "del tronco, distacco, stabilizzazione in piedi, discesa) sulla "
    "base di soglie cinematiche definite a priori.")
body(doc,
    "(4) Estrazione degli indicatori biomeccanici (KPI). Per ogni "
    "ripetizione saranno calcolati 18 KPI raggruppati in 5 categorie "
    "anatomiche: tronco (angolo di flessione anteriore, simmetria di "
    "inclinazione), ginocchio (angolo di flessione, valgismo/varismo, "
    "velocità angolare massima in estensione), anca e velocità "
    "(velocità verticale del bacino al distacco, angolo di flessione "
    "d'anca), simmetria e equilibrio (differenza laterale di "
    "carico-proxy, escursione laterale del centro di massa), durata di "
    "fase (tempi di ciascuna delle cinque fasi e tempo totale per "
    "ripetizione). Per ogni partecipante e per ogni KPI saranno "
    "calcolati i parametri descrittivi (media, deviazione standard, "
    "minimo, massimo) sull'intera serie di ripetizioni.")
body(doc,
    "(5) Personalizzazione mediante baseline z-score individuali. "
    "Caratteristica distintiva della pipeline ARISE è l'adozione di un "
    "approccio di valutazione personalizzata per singolo partecipante. "
    "Per ciascun soggetto, le prime N ripetizioni ritenute conformi "
    "(N ≥ 5, parametro configurabile) vengono utilizzate per "
    "costruire una baseline individuale dei principali KPI: per ogni "
    "indicatore vengono calcolate media e deviazione standard "
    "personali. Ogni nuova ripetizione viene successivamente valutata "
    "rispetto a questa baseline calcolandone il punteggio z (z-score); "
    "una deviazione |z| > 2.0, corrispondente al 97.5° percentile "
    "della distribuzione gaussiana individuale, segnala un'esecuzione "
    "atipica per quel particolare soggetto. Quando il numero di "
    "ripetizioni di calibrazione è insufficiente (cold-start), si "
    "applicano soglie cinematiche predefinite (rule-based) ricavate "
    "dalla letteratura. Questo approccio dual-mode (regole + z-score "
    "personalizzato) rende la valutazione robusta alla variabilità "
    "antropometrica e biomeccanica del singolo soggetto e indipendente "
    "da soglie hardcoded uguali per tutti.")
body(doc,
    "(6) Modellazione mediante machine learning. In parallelo al "
    "canale rule-based + z-score, la serie temporale dei keypoint per "
    "ogni ripetizione verrà data in ingresso a una rete neurale "
    "ricorrente (Long Short-Term Memory, LSTM) addestrata a codificare "
    "il pattern dinamico del Sit-to-Stand in una rappresentazione "
    "vettoriale compatta (embedding). Tale rappresentazione sarà "
    "successivamente utilizzata da un classificatore (es. random "
    "forest o rete densa) addestrato a produrre un punteggio di "
    "normalità per ripetizione e a discriminare pattern tipici da "
    "pattern atipici. I dati del presente studio costituiranno la "
    "popolazione di riferimento sana sulla quale i modelli "
    "apprenderanno la baseline di normalità. L'addestramento sarà "
    "condotto con validazione incrociata leave-one-subject-out, in "
    "modo che ciascun partecipante sia escluso a turno dal training "
    "set e utilizzato come test set indipendente, garantendo una "
    "stima realistica della capacità di generalizzazione su nuovi "
    "soggetti. Le metriche di performance includeranno accuratezza, "
    "F1-score, sensibilità, specificità e area sotto la curva ROC "
    "(AUC). Nessun dato individuale verrà trasferito al di fuori del "
    "perimetro del progetto, e nessun modello addestrato sarà reso "
    "pubblico in modo da permettere la re-identificazione dei "
    "partecipanti.")
body(doc,
    "(7) Analisi statistica. (a) Accordo con la strumentazione di "
    "riferimento (sotto-coorte facoltativa): per ogni KPI saranno "
    "calcolati Bland-Altman plot (bias e limiti di accordo al 95%), "
    "Concordance Correlation Coefficient di Lin, Intraclass Correlation "
    "Coefficient (ICC, modello 2,1) ed errore assoluto medio (MAE). "
    "(b) Caratterizzazione della popolazione sana: per ogni KPI e per "
    "ciascuna fascia d'età (adulti, anziani) saranno riportati media, "
    "deviazione standard, mediana e percentili 5°/25°/75°/95°. La "
    "normalità delle distribuzioni sarà verificata con il test di "
    "Shapiro-Wilk. (c) Confronto tra fasce d'età: a seconda della "
    "distribuzione si applicherà t-test indipendente o Mann-Whitney U "
    "con correzione per confronti multipli (Benjamini-Hochberg). (d) "
    "Variabilità inter e intra-soggetto: sarà quantificata mediante "
    "coefficiente di variazione (CV) e ICC per misure ripetute. "
    "(e) Confronto fra i due canali di analisi (rule-based/z-score vs "
    "LSTM): concordanza valutata con coefficiente kappa di Cohen e "
    "matrice di confusione. L'elaborazione e l'addestramento dei "
    "modelli saranno condotti in Python (NumPy, SciPy, pandas, "
    "PyTorch, scikit-learn) e R; tutto il codice sarà versionato e "
    "l'analisi sarà completamente riproducibile.")

label(doc, "Risultati attesi:")
body(doc,
    "Costituzione di un corpus di dati biomeccanici di riferimento "
    "sul Sit-to-Stand in soggetti sani nelle diverse fasce d'età; "
    "verifica della concordanza tra i dati raccolti dai diversi "
    "sistemi; pubblicazioni scientifiche nel campo della biomeccanica "
    "del movimento e dell'analisi del movimento senza marker.")

label(doc, "Riferimenti bibliografici principali:")
body(doc,
    "[1] Bohannon RW. Sit-to-stand test for measuring performance of "
    "lower extremity muscles. Percept Mot Skills 1995;80:163-6.\n"
    "[2] Jones CJ, Rikli RE, Beam WC. A 30-s chair-stand test as a "
    "measure of lower body strength in community-residing older adults. "
    "Res Q Exerc Sport 1999;70:113-9.\n"
    "[3] Lugade V, Lin V, Chou L-S. Validity of using markerless motion "
    "capture for biomechanical assessment. Gait Posture 2022.\n"
    "[4] Teresi JA, Yu X, Stewart AL, Hays RD. Guidelines for Designing "
    "and Evaluating Feasibility Pilot Studies. Med Care 2022;60:95-103.\n"
    "[5] Janssen WGM, Bussmann HBJ, Stam HJ. Determinants of the "
    "sit-to-stand movement: a review. Phys Ther 2002;82:866-79.\n"
    "[6] Millor N, Lecumberri P, Gomez M, Martinez-Ramirez A, Izquierdo M. "
    "Kinematic parameters to evaluate functional performance of "
    "sit-to-stand and stand-to-sit transitions. IEEE Trans Neural Syst "
    "Rehabil Eng 2014;22:926-34.\n"
    "[7] Hertzog MA. Considerations in determining sample size for pilot "
    "studies. Res Nurs Health 2008;31:180-91.\n"
    "[8] Julious SA. Sample size of 12 per group rule of thumb for a "
    "pilot study. Pharm Stat 2005;4:287-91.\n"
    "[9] Csuka M, McCarty DJ. Simple method for measurement of lower "
    "extremity muscle strength. Am J Med 1985;78:77-81.")

label(doc, "Eventuali enti finanziatori esterni o Sponsor:")
body(doc,
    "Lo studio si inserisce nel progetto ARISE, finanziato dal Fondo "
    "Europeo di Sviluppo Regionale (FESR/ERDF), programma HealthTech "
    "regionale della Regione Liguria. Il partner tecnico-industriale "
    "del progetto è Innovina S.r.l., che fornisce il software di "
    "acquisizione e analisi video utilizzato nello studio.")

field(doc, "Data prevista di inizio della ricerca",
      "Settembre 2026, comunque successivi al ricevimento del parere "
      "favorevole CERA.")
field(doc, "Durata prevista della ricerca",
      "Circa 12 mesi complessivi (acquisizione, analisi, redazione dei "
      "risultati). Coinvolgimento individuale del singolo partecipante: "
      "2 sessioni di 10-15 minuti ciascuna nell'arco di una settimana.")

# 3. Partecipanti
spacer(doc)
heading1(doc, "3. INFORMAZIONI RELATIVE AI PARTECIPANTI COINVOLTI NELLA RICERCA")

label(doc, "Campione")
field(doc, "Numero",
      "50 partecipanti totali, suddivisi in differenti fasce d'età "
      "(adulti sani 18-64 anni e ≥65 anni). Tale numerosità è "
      "supportata da indicazioni metodologiche [4], [7], [8] e risulta "
      "in linea con gli studi di riferimento per la caratterizzazione "
      "biomeccanica del Sit-to-Stand [1], [9].")

label(doc, "Sesso", bold=False)
checks(doc, [
    (False, "maschi"),
    (False, "femmine"),
    (True,  "maschi e femmine"),
])

label(doc, "Età", bold=False)
checks(doc, [
    (True,  "adulti (≥18 anni)"),
    (False, "minori"),
    (False, "adulti e minori"),
])

label(doc, "Coinvolge persone con disabilità?")
checks(doc, [
    (False, "Sì"),
    (True,  "No"),
])

spacer(doc)
label(doc, "Nello specifico, quali tipologie di soggetti prenderanno parte allo studio?")
checks(doc, [
    (False, "Studenti"),
    (True,  "Adulti sani (>18 anni, <65 anni)"),
    (False, "Bambini e ragazzi di età inferiore a 18 anni"),
    (True,  "Anziani sani (>65 anni)"),
    (False, "Soggetti di madrelingua non italiana"),
    (False, "Soggetti con deficit cognitivo/mentale"),
    (False, "Altre persone la cui capacità di esprimere consenso possa "
            "essere compromessa"),
    (False, "Soggetti istituzionalizzati"),
    (False, "Pazienti e/o clienti segnalati da medici, psicologi o altre "
            "categorie di professionisti"),
    (False, "Non è possibile determinare la tipologia di soggetti"),
])

label(doc, "Descrizione del gruppo dei soggetti partecipanti:")
body(doc,
    "I/le partecipanti, di qualsiasi sesso e genere, saranno adulti "
    "(>18 anni) in buono stato di salute generale, senza patologie "
    "note del sistema muscolo-scheletrico, cardiovascolare o "
    "neurologico tali da controindicare l'esecuzione ripetuta del "
    "Sit-to-Stand. I/le partecipanti devono essere in grado di "
    "alzarsi e sedersi autonomamente da una sedia standard senza "
    "assistenza fisica.")

label(doc, "Modalità di reclutamento dei soggetti partecipanti alla ricerca:")
body(doc,
    "La partecipazione alla sperimentazione sarà volontaria e su base "
    "propositiva. L'utilizzo di passaparola tra persone sarà concesso.")

label(doc, "Criteri di inclusione/esclusione dei partecipanti alla ricerca:")
body(doc,
    "Criteri di inclusione: età ≥18 anni; buono stato di salute generale; "
    "capacità di eseguire il Sit-to-Stand senza assistenza fisica esterna; "
    "capacità di comprendere le istruzioni in lingua italiana; firma del "
    "consenso informato e presa visione dell'informativa privacy.")
body(doc,
    "Criteri di esclusione: controindicazione medica all'esecuzione "
    "ripetuta del Sit-to-Stand (es. patologie cardiache scompensate, "
    "ipotensione ortostatica grave, vertigini significative); lesione "
    "acuta o intervento chirurgico recente agli arti inferiori o al "
    "tronco (entro 6 settimane); patologia neurologica o muscolo-"
    "scheletrica che comprometta l'esecuzione autonoma del movimento; "
    "compromissione cognitiva; partecipazione concorrente ad altra "
    "sperimentazione interventistica.")

label(doc,
    "È possibile che alcuni dei soggetti si trovino in una posizione di "
    "dipendenza nei confronti del ricercatore o dei suoi collaboratori, "
    "tale per cui l'espressione del consenso non sia del tutto libera?")
body(doc,
    "La partecipazione al progetto di ricerca è del tutto volontaria; "
    "pertanto, non sarà fatto obbligo a nessuno di prendervi parte. "
    "Coloro che desiderano partecipare potranno farlo in via del tutto "
    "autonoma, contattando uno dei membri del team di ricerca via "
    "telefonica o via e-mail. Inoltre, essendo i dati raccolti in "
    "forma pseudonimizzata, non vi sarà modo di sapere chi ha preso "
    "parte alla sperimentazione.")

# 4. Rischio
spacer(doc)
heading1(doc, "4. RISCHIO E GESTIONE DEL RISCHIO")

label(doc, "Metodi di indagine")
checks(doc, [
    (False, "Somministrazione di test standardizzati"),
    (False, "Raccolta dati tramite colloquio clinico"),
    (False, "Raccolta dati tramite interviste"),
    (False, "Raccolta dati tramite questionari"),
    (False, "Raccolta dati archivi"),
    (True,  "Attività pratiche (esecuzione del Sit-to-Stand secondo "
            "protocolli standardizzati 5xSTS e 30sCST)"),
    (True,  "Fotografie e/o Videoregistrazioni: registrazione video con "
            "telecamere RGB e/o RGB-D necessaria all'analisi "
            "biomeccanica del movimento."),
])

label(doc, "Nello specifico, la ricerca prevede:")
checks(doc, [
    (False, "utilizzo di questionari"),
    (False, "interviste strutturate o semi-strutturate"),
    (False, "narrazioni autobiografiche"),
    (False, "raccolta di diari"),
    (False, "osservazione del comportamento dei soggetti a loro insaputa"),
    (True,  "osservazione del comportamento dei soggetti (con piena "
            "consapevolezza e consenso del partecipante)"),
    (False, "registrazioni audio dei soggetti"),
    (True,  "registrazioni video dei soggetti"),
    (True,  "somministrazione di stimoli, compiti o procedure e "
            "registrazione di risposte comportamentali, opinioni o "
            "giudizi (esecuzione del Sit-to-Stand)"),
    (False, "somministrazione di stimoli o procedure fastidiosi/stressanti"),
    (False, "registrazione di movimenti oculari"),
    (False, "immersione in ambienti di realtà virtuale"),
    (False, "somministrazione di test o questionari attraverso internet"),
    (False, "utilizzo di test neuropsicologici"),
    (False, "comportamenti che potrebbero diminuire l'autostima"),
    (False, "procedure di inganno dei soggetti"),
    (True,  "Altro: acquisizione cinematica tridimensionale del movimento "
            "mediante pose estimation senza marker da video RGB e/o RGB-D."),
])

label(doc, "Rischi per i soggetti")
checks(doc, [
    (False, "Nessun rischio"),
    (False, "Intrusività degli strumenti"),
    (True,  "Affaticamento (correlato all'esecuzione ripetuta del Sit-to-Stand)"),
    (False, "Dolore fisico"),
    (False, "Possibile induzione di sentimenti di frustrazione o autosvalutazione"),
    (False, "Forte tensione emotiva"),
    (False, "Effetti negativi sullo stato emotivo"),
])

label(doc,
    "In caso di rischi, precisare natura e intensità, giustificare la "
    "ricerca alla luce di danni/benefici, specificare precauzioni:")
body(doc,
    "La partecipazione presenta rischi minimi e sovrapponibili a "
    "quelli di una qualsiasi attività motoria leggera. Il lieve "
    "affaticamento muscolare correlato alle ripetizioni è mitigato "
    "dal numero contenuto di ripetizioni per sessione e dalla "
    "possibilità per il partecipante di effettuare pause o "
    "interrompere su richiesta. Gli esercizi saranno eseguiti sotto "
    "supervisione continua di un operatore di ricerca presente "
    "accanto al partecipante che ne assicurerà la sicurezza e un "
    "continuo monitoraggio del livello di affaticamento. Non si "
    "prevedono rischi di natura psicologica. Il partecipante può "
    "interrompere la propria partecipazione in qualsiasi momento "
    "senza fornire alcuna motivazione e senza alcuna conseguenza; in "
    "tal caso i dati raccolti non verranno conservati.")
body(doc,
    "Rischi per la privacy: le registrazioni video sono "
    "pseudonimizzate alla sorgente mediante codice identificativo "
    "randomico generato automaticamente dal software di acquisizione, "
    "in modo che il/la partecipante non sia direttamente identificabile "
    "a partire dai dati raccolti. Le modalità di archiviazione e di "
    "controllo degli accessi sono descritte nella Sezione 6.")

label(doc, "Si prevede che vi possano essere benefici per chi prende parte alla ricerca?")
body(doc,
    "Per la partecipazione alla ricerca non sono previsti benefici "
    "materiali, economici o di altro tipo.")

# 5. Consenso
spacer(doc)
heading1(doc, "5. INFORMAZIONE, CONSENSO E RESTITUZIONE")

label(doc, "Informazione al partecipante sugli scopi della ricerca")
checks(doc, [
    (True,  "Prima dell'intervento"),
    (False, "Dopo l'intervento"),
    (False, "Durante l'intervento"),
])

label(doc, "Come verranno diffusi le informazioni e l'invito a partecipare:")
body(doc,
    "Gli inviti saranno diffusi tramite e-mail, mailing list "
    "dipartimentali dell'Università di Genova, locandine affisse presso "
    "il Dipartimento DINOGMI, e passaparola tra persone. Il materiale "
    "informativo (Nota Informativa, Informativa Privacy, Modulo di "
    "Consenso) è allegato alla presente richiesta.")

label(doc, "Consenso (modulo allegato)")
label(doc, "A chi è richiesto il consenso informato?", bold=False)
checks(doc, [
    (True,  "Soggetti partecipanti"),
    (False, "Genitori (in caso di minori): non applicabile"),
    (False, "Responsabile dell'Istituzione dove si svolge la raccolta "
            "dati (non necessario, i Laboratori del Dipartimento DINOGMI "
            "sono sede del gruppo di ricerca)"),
])

label(doc, "Modalità di gestione di dubbi e richieste di precisazioni:")
body(doc,
    "Ai/alle partecipanti allo studio verrà mostrata l'informativa "
    "dello studio su cui saranno specificati gli obiettivi e i "
    "responsabili dello studio con i relativi indirizzi e-mail e numeri "
    "di telefono. La stessa informativa sarà rilasciata in duplice copia "
    "ai/alle partecipanti che in caso di dubbi o necessità di maggiori "
    "informazioni potranno riconsultare ed eventualmente contattare i "
    "responsabili prima, durante e dopo la partecipazione allo studio. "
    "I/le partecipanti potranno scegliere di abbandonare il test anche "
    "in itinere senza dover fornire alcuna spiegazione.")

label(doc, "Sintesi del percorso informativo previsto:")
body(doc,
    "Le persone che decideranno di prendere parte allo studio verranno "
    "adeguatamente informate sia verbalmente che tramite la lettura "
    "dei documenti dello studio (Nota Informativa, Informativa "
    "Privacy). Successivamente verrà richiesto di firmare il consenso "
    "informato e il consenso al trattamento dei dati. In caso di "
    "consenso negato non verrà emesso alcun giudizio. Preliminarmente "
    "alla sperimentazione verranno raccolti dati socio-demografici "
    "generici (età, altezza, peso, sesso assegnato alla nascita, "
    "genere). Al termine, i/le partecipanti verranno informati della "
    "restituzione dei risultati tramite e-mail. I risultati aggregati "
    "anonimi potranno comunque essere reperiti autonomamente da "
    "chiunque attraverso la pubblicazione dello studio.")

label(doc, "Restituzione")
body(doc,
    "Al termine di ogni sessione, su richiesta, il/la partecipante "
    "potrà ricevere un breve feedback verbale dal personale di ricerca "
    "sulla propria performance biomeccanica osservata durante "
    "l'esecuzione del task. I risultati aggregati e anonimi dello "
    "studio saranno resi disponibili attraverso le pubblicazioni "
    "scientifiche derivanti dalla ricerca e potranno essere richiesti "
    "ai ricercatori coinvolti.")

# 6. Dati
spacer(doc)
heading1(doc, "6. ANONIMATO / RISERVATEZZA DEI DATI PERSONALI E "
              "CONSERVAZIONE / SICUREZZA DEI RISULTATI DELLA RICERCA")

label(doc, "Tipo dati per trattamento")
checks(doc, [
    (False, "Dati raccolti in forma anonima"),
    (True,  "Dati raccolti in forma riservata (pseudonimizzata)"),
])

label(doc, "Altre informazioni circa l'anonimato dei partecipanti")
label(doc, "Come verrà garantito ai partecipanti l'anonimato "
           "(ad es., utilizzo di codici di identificazione)?", bold=False)
body(doc,
    "Per ogni partecipante verrà utilizzato un codice di identificazione "
    "caratterizzato da una sequenza numerica randomica (es. sbj4112, "
    "sbj1252) che non permetterà di risalire ai dati personali del/della "
    "partecipante. I dati saranno raccolti ed elaborati nel pieno "
    "rispetto delle normative vigenti sulla privacy (GDPR Regolamento "
    "Europeo UE 2016/679; D.Lgs. 10 agosto 2018, n. 101), rendendo "
    "pseudonimizzati i dati e conservando il materiale raccolto su "
    "server protetto dell'Università.")

label(doc, "Nel caso fosse necessario conservare i dati identificativi "
           "dei partecipanti, specificarne i motivi e le modalità con cui "
           "i soggetti ne sono informati:")
body(doc,
    "Nil. La mappatura codice / identità, ove conservata per le "
    "finalità di esercizio dei diritti GDPR del/della partecipante, "
    "sarà mantenuta in archivio cifrato presso il Responsabile della "
    "ricerca.")

label(doc, "Chi avrà accesso ai dati raccolti e ai risultati, ancorché "
           "intermedi, della ricerca?")
body(doc,
    "L'accesso ai dati identificativi sarà limitato al Responsabile "
    "della ricerca. L'accesso ai dati pseudonimizzati e ai risultati "
    "preliminari sarà garantito esclusivamente ai componenti del "
    "gruppo di ricerca (Prof. Marco Testa e dottorando) e a Innovina "
    "S.r.l. nella sua qualità di Responsabile esterno del trattamento "
    "ex Art. 28 GDPR, limitatamente alle finalità di elaborazione "
    "tecnica, addestramento dei modelli e analisi previste dal "
    "progetto e regolate dal Data Processing Agreement. Nessun altro "
    "ente terzo potrà accedere ai dati raccolti. I risultati "
    "aggregati e anonimi saranno resi pubblici tramite pubblicazione "
    "scientifica.")

label(doc, "Indicare le modalità di conservazione dei dati sensibili "
           "(responsabile della corretta conservazione e luogo dove "
           "verranno conservati):")
body(doc,
    "Il trattamento dei dati personali, sebbene riservati, sarà "
    "effettuato con modalità automatizzata. Il materiale sensibile "
    "sarà custodito dal Responsabile della ricerca e protetto da "
    "password in apposito archivio. In particolare, i dati saranno "
    "archiviati sul server protetto del Dipartimento DINOGMI, "
    "utilizzato per effettuare le acquisizioni, e che resterà sempre "
    "all'interno della sede del dipartimento. Solo i ricercatori "
    "coinvolti nel progetto potranno accedervi. I dati saranno "
    "conservati per 3 anni dopo il termine dello studio. I dati "
    "sensibili verranno trattati e conservati secondo la normativa "
    "vigente del Regolamento UE n. 2016/679.")

spacer(doc, 2)
body(doc,
    "In relazione allo svolgimento del progetto dal titolo "
    "\"ARISE: A feasibility study on markerless video-based biomechanical "
    "assessment of the Sit-to-Stand task in healthy adults and older "
    "adults\".")
spacer(doc)
p = doc.add_paragraph()
r = p.add_run("Si chiede l'espressione del parere da parte del "
              "CERA, Università degli Studi di Genova."); r.bold = True; r.font.size = Pt(11)
body(doc,
    "Il responsabile della ricerca dichiara di conoscere il "
    "Regolamento del CERA, Università degli Studi di Genova.")

spacer(doc, 3)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prof. Marco Testa PT, PhD"); r.bold = True; r.font.size = Pt(11)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Delegato del Rettore per il Campus Universitario di Savona"); r.font.size = Pt(11)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Presidente del Master in Riabilitazione dei Disordini Muscoloscheletrici"); r.font.size = Pt(11)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Università di Genova, Campus of Savona"); r.font.size = Pt(11)

spacer(doc, 2)
heading1(doc, "ELENCO DEGLI ALLEGATI")
attachments = [
    "Allegato 1: Nota Informativa per il partecipante",
    "Allegato 2: Informativa Privacy ex Art. 13 GDPR",
    "Allegato 3: Modulo di Consenso Informato",
    "Allegato 4: Data Processing Agreement (DPA) tra UNIGE e Innovina "
                "S.r.l. ex Art. 28 GDPR (in corso di formalizzazione)",
    "Allegato 5: Dichiarazione sul Conflitto di Interessi",
]
for a in attachments:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run("• " + a); r.font.size = Pt(11)

doc.save(OUT)
print(f"Wrote {OUT}")
