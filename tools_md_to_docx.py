"""Convert ARISE markdown deliverables to professional .docx."""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
LINK_RE = re.compile(r"\[(.+?)\]\((.+?)\)")
INLINE_CODE_RE = re.compile(r"`(.+?)`")


def add_runs_with_formatting(paragraph, text):
    """Parse a line of markdown inline formatting and add runs."""
    # Strategy: tokenize on **bold**, *italic*, `code`, [link](url)
    # Simple sequential approach: handle bold first, then italic in remainder.
    # For deliverable docs this is sufficient — no nested formatting.
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("["):
            lm = LINK_RE.match(token)
            run = paragraph.add_run(lm.group(1))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_table(doc, rows):
    """rows = list of list of cell strings; first row is header unless empty."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    # Detect "header-less" cover-style tables (first row entirely empty)
    has_header = any(cell.strip() for cell in rows[0])
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Table Grid"  # plain black borders, no fill
    if not has_header:
        rows = rows[1:]
        tbl._element.remove(tbl.rows[0]._tr)
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs_with_formatting(p, cell_text.strip())
            if has_header and i == 0:
                # Bold header text only, no background fill
                for run in p.runs:
                    run.bold = True
            elif not has_header and j == 0:
                # Cover-style: bold the left-column label
                for run in p.runs:
                    run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def parse_table_block(lines, i):
    """Parse a contiguous markdown table starting at lines[i]. Returns (rows, next_i)."""
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        # skip separator row e.g. |---|---|
        if re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line):
            i += 1
            continue
        # split, drop empty first/last from leading/trailing |
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        rows.append(cells)
        i += 1
    return rows, i


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # Page setup: A4, reasonable margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Footer with page numbers
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Heading styles: plain black, conservative sizes (like a manager would use)
    for lvl, size in [(1, 16), (2, 13), (3, 12), (4, 11)]:
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        s.font.bold = True

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule → skip (visual separator only)
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            # Strip surrounding ** if present (some titles wrap in bold)
            heading_text = re.sub(r"^\*\*(.+)\*\*$", r"\1", heading_text)
            p = doc.add_heading(level=min(level, 4))
            add_runs_with_formatting(p, heading_text)
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            rows, i = parse_table_block(lines, i)
            add_table(doc, rows)
            continue

        # Bullet list
        if re.match(r"^[-*+]\s+", stripped):
            while i < n and re.match(r"^[-*+]\s+", lines[i].strip()):
                item = re.sub(r"^[-*+]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_runs_with_formatting(p, item)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_runs_with_formatting(p, item)
                i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            content = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            r = p.add_run(content)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing ```
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            continue

        # Regular paragraph — combine continuation lines until blank
        para_lines = [line]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#|\||[-*+]\s|\d+\.\s|>\s|```|-{3,})", lines[i].strip()
        ):
            para_lines.append(lines[i])
            i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        add_runs_with_formatting(p, para_text)

    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(".docx")
    convert(src, dst)
