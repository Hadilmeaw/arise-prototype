"""ARISE — Informativa Privacy (paragraph-style format), matching the
UNIGE/DINOGMI template "InformativaPrivacy.docx". Includes Sezione A
(GDPR Art. 13 informativa) and Sezione B (consent forms for adults, minors
capable of understanding, and minors with legal representative).
The minor sections are retained from the template; ARISE does not enrol
minors but the sections are kept as required by the standard form.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2")
OUT  = ROOT / "ARISE_CERA_Informativa_Privacy.docx"

STUDY_TITLE = (
    "ARISE — A feasibility study on markerless video-based biomechanical "
    "assessment of the Sit-to-Stand task in healthy adults and older adults"
)

def hcenter(doc, t, *, bold=False, size=11):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)

def section(doc, t, *, size=12):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = True; r.font.size = Pt(size)

def label(doc, t):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = True; r.font.size = Pt(11)

def body(doc, t, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(t); r.italic = italic; r.font.size = Pt(size)

def numpoint(doc, n, t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"{n}. {t}"); r.font.size = Pt(11)

def footnote(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(t); r.italic = True; r.font.size = Pt(9)

def checkbox(doc, t):
    p = doc.add_paragraph()
    r = p.add_run("☐  " + t); r.font.size = Pt(11)

def spacer(doc, n=1):
    for _ in range(n): doc.add_paragraph()

doc = Document()
s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)

# ===================== HEADER =====================
hcenter(doc, "UNIVERSITÀ DEGLI STUDI DI GENOVA", bold=True, size=14)
hcenter(doc, "Dipartimento di Neuroscienze, Riabilitazione, Oftalmologia, "
             "Genetica e Scienze Materno-Infantili")
spacer(doc)

# ===================== SEZIONE A =====================
section(doc, "Sezione A)")
spacer(doc)
hcenter(doc,
    "INFORMATIVA TRATTAMENTO DATI PERSONALI NELL'AMBITO DEI PROGETTI DI RICERCA",
    bold=True, size=11)
hcenter(doc,
    "RESA AI SENSI DELL'ART. 13 DEL REGOLAMENTO UE N. 2016/679",
    bold=True, size=11)
spacer(doc)

body(doc, "Gentile Partecipante,")
body(doc,
    "desideriamo informarLa che il GDPR (Regolamento UE n. 679/2016) "
    "prevede la tutela delle persone rispetto al trattamento dei dati "
    "personali.")
body(doc,
    "Secondo la normativa indicata, tale trattamento sarà improntato ai "
    "principi di liceità, correttezza e trasparenza, adeguatezza, "
    "pertinenza e limitazione, esattezza e aggiornamento, non "
    "eccedenza¹ e responsabilizzazione.")

spacer(doc)
label(doc, "Informazioni sul trattamento dei suoi dati.")

numpoint(doc, 1,
    f"I dati da Lei forniti verranno trattati per finalità di ricerca "
    f"scientifica, nell'ambito del progetto di ricerca dal titolo "
    f"\"{STUDY_TITLE}\".")
body(doc,
    "La ricerca è finalizzata a (1) raccogliere dati biomeccanici di "
    "riferimento sull'esecuzione del Sit-to-Stand in una popolazione di "
    "adulti (>18 anni); (2) confrontare le misure ottenute mediante "
    "analisi video senza marker con quelle ottenute da sistemi di "
    "analisi cinematica e cinetica del movimento gold-standard validati "
    "con metodologie non invasive; (3) verificare la fattibilità "
    "dell'approccio video senza marker per la valutazione biomeccanica "
    "del movimento.")
body(doc,
    "A tal fine ai/alle partecipanti sarà chiesto di fornire alcune "
    "informazioni sociodemografiche (età, altezza, peso, sesso, genere) "
    "e anamnestiche generiche sullo stato di salute, i propri dati di "
    "contatto (nome, cognome, recapito telefonico, e-mail) e di "
    "partecipare alla sperimentazione presso i Laboratori di ricerca "
    "del Dipartimento DINOGMI (Campus Universitario di Savona), "
    "eseguendo alcune ripetizioni del task Sit-to-Stand registrate "
    "mediante telecamere RGB e/o RGB-D.")
body(doc,
    "I dati saranno raccolti in modo manuale ed automatico, tramite "
    "modalità che ne garantiscano la sicurezza, la riservatezza e la "
    "pseudonimizzazione.")

numpoint(doc, 2,
    "Il trattamento dei dati personali sarà effettuato con modalità "
    "automatizzata. Il titolare del trattamento metterà in atto misure "
    "tecniche e organizzative adeguate volte ad attuare in modo "
    "efficace i principi di protezione dei dati e a tutelare i diritti "
    "degli interessati.")

numpoint(doc, 3,
    "Il conferimento dei dati è facoltativo, cioè non discende da un "
    "obbligo normativo, e l'eventuale rifiuto di fornire tali dati non "
    "ha conseguenze. Il conferimento dei dati per le finalità di cui al "
    "punto 1 è indispensabile allo svolgimento dello studio. Il rifiuto "
    "di conferirli non consentirà all'Interessato di partecipare allo "
    "studio in parola.")

numpoint(doc, 4,
    "Il conferimento dei dati di contatto (nome, cognome, recapito "
    "telefonico, e-mail) è facoltativo, cioè non discende da un obbligo "
    "normativo, ma è necessario per consentire l'esercizio dei diritti "
    "GDPR del/della partecipante (in particolare il diritto di accesso, "
    "di rettifica, di cancellazione e di restituzione dei risultati "
    "personali), nonché per consentire al Titolare di ricontattare "
    "l'Interessato affinché possa esprimere, se lo riterrà, un nuovo "
    "specifico consenso per una ulteriore ricerca. Il mancato "
    "conferimento dei dati per tali finalità avrà come unica "
    "conseguenza l'impossibilità di realizzare quanto da ultimo "
    "descritto.")

numpoint(doc, 5,
    "Il trattamento potrebbe riguardare anche dati personali rientranti "
    "nel novero dei dati particolari, tra i quali dati idonei a "
    "rivelare lo stato di salute. I dati sanitari saranno trattati per "
    "esclusiva finalità della ricerca.")

numpoint(doc, 6,
    "I dati identificativi diretti (nome, cognome, recapito telefonico, "
    "e-mail) non saranno comunicati ad alcun soggetto terzo. I dati "
    "pseudonimizzati (video e indicatori cinematici etichettati con "
    "codice randomico, es. sbj4112) saranno comunicati esclusivamente "
    "a Innovina S.r.l. nella sua qualità di Responsabile esterno del "
    "trattamento ai sensi dell'Art. 28 GDPR, vincolata da apposito Data "
    "Processing Agreement (DPA) attualmente in corso di formalizzazione, "
    "ai soli fini dell'elaborazione "
    "tecnica e dell'addestramento dei modelli previsti dal progetto. "
    "I risultati saranno oggetto di diffusione esclusivamente in forma "
    "anonima ed aggregata.")

numpoint(doc, 7,
    "I dati saranno conservati per 3 anni dopo il termine dello studio. "
    "I file ed i documenti contenenti i dati dello studio saranno "
    "conservati ed accessibili solo tramite password, su server "
    "protetto del Dipartimento DINOGMI. Tutto il materiale sarà "
    "mantenuto riservato e sarà sottoposto al trattamento nel pieno "
    "rispetto del D.Lgs. 196/03 e del Regolamento UE n. 2016/679. "
    "L'ultima revisione della dichiarazione di Helsinki nonché la "
    "dichiarazione di Oviedo sono la base per la conduzione etica "
    "dello studio.")

spacer(doc)
footnote(doc,
    "¹ La non eccedenza si sostanzia in quello che il Codice Privacy "
    "definisce come principio di necessità: occorre ridurre al minimo "
    "l'utilizzo di dati personali e identificativi, in modo da "
    "escluderne il trattamento, qualora le finalità perseguite siano "
    "raggiungibili anche mediante dati anonimi o opportune modalità "
    "che permettano l'identificazione dell'interessato solo nei casi "
    "di necessità (cd. Minimizzazione).")

spacer(doc)
label(doc, "Titolare, responsabile, autorizzati al trattamento.")

numpoint(doc, 1,
    "Il Titolare del trattamento è l'Università di Genova con sede in "
    "Genova, via Balbi 5, nella persona del Rettore pro tempore "
    "(dati di contatto reperibili alla pagina "
    "https://intranet.unige.it/privacy).")

numpoint(doc, 2,
    "Il Responsabile del trattamento è prof. Mohammad Maghnie, "
    "Direttore del Dipartimento di Neuroscienze, Riabilitazione, "
    "Oftalmologia, Genetica e Scienze Materno-Infantili dell'Università "
    "di Genova (email: mohamad.maghnie@unige.it).")

numpoint(doc, 3,
    "Il Coordinatore del progetto è Prof. Marco Testa "
    "(email: marco.testa@unige.it).")

numpoint(doc, 4,
    "Le persone autorizzate al trattamento dei dati identificativi "
    "diretti del partecipante (nome, cognome, recapito telefonico, "
    "e-mail) sono esclusivamente il Coordinatore del progetto, Prof. "
    "Marco Testa. Le persone autorizzate al trattamento dei dati "
    "pseudonimizzati sono Prof. Marco Testa e un dottorando del "
    "gruppo di ricerca. Innovina S.r.l., designata Responsabile "
    "esterno del trattamento ai sensi dell'Art. 28 GDPR e vincolata "
    "da apposito Data Processing Agreement (DPA) attualmente in corso "
    "di formalizzazione, potrà ricevere esclusivamente i dati delle "
    "acquisizioni in forma pseudonimizzata ai fini dell'elaborazione "
    "tecnica e dell'addestramento dei modelli previsti dal progetto.")

numpoint(doc, 5,
    "Presso il Titolare del trattamento è presente il Responsabile "
    "della protezione dei dati, nominato ai sensi dell'art. 37 del "
    "Regolamento UE 2016/679. Il responsabile della protezione dei "
    "dati può essere contattato all'indirizzo dpo@unige.it.")

spacer(doc)
label(doc, "Contatti e segnalazioni.")
body(doc,
    "Per informazioni circa il progetto o eventuali segnalazioni è "
    "possibile contattare i ricercatori coinvolti nel progetto:")
contacts = [
    ("Prof. Marco Testa", "+39 331 2611548", "marco.testa@unige.it"),
    ("Paolo Pavani (Innovina S.r.l.)",   "[telefono da inserire]", "[email da inserire]"),
    ("Hadil Sahraoui (Innovina S.r.l.)", "[telefono da inserire]", "[email da inserire]"),
]
for name, phone, email in contacts:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f"{name}   ·   Cellulare: {phone}   ·   {email}")
    r.font.size = Pt(11)

spacer(doc)
body(doc,
    "In ogni momento potrà esercitare i Suoi diritti di informazione e "
    "accesso, di rettifica e cancellazione, di limitazione e di "
    "opposizione al trattamento, di portabilità dei dati personali "
    "(di cui alla sezione 2, 3 e 4 del capo III del Regolamento UE n. "
    "679/2016).")

doc.add_page_break()

# ===================== SEZIONE B — ADULTS =====================
hcenter(doc, "UNIVERSITÀ DEGLI STUDI DI GENOVA", bold=True, size=14)
hcenter(doc, "Dipartimento di Neuroscienze, Riabilitazione, Oftalmologia, "
             "Genetica e Scienze Materno-Infantili")
spacer(doc)

section(doc, "Sezione B")
spacer(doc)
hcenter(doc,
    "CONSENSO INFORMATO PER IL TRATTAMENTO DEI DATI PERSONALI "
    "NELL'AMBITO DEI PROGETTI DI RICERCA",
    bold=True, size=11)
hcenter(doc,
    "da compilare nel caso di raccolta di dati personali particolari "
    "(ex artt. 9 e 10 del Reg. UE n. 2016/679) di soggetti maggiorenni",
    bold=True, size=10)
spacer(doc)

body(doc, "Il/la sottoscritto/a ………………………………………………………………………")
spacer(doc)

checkbox(doc, "DICHIARA di aver preso visione dell'Informativa per il "
              "trattamento dei dati personali di cui alla sezione A")
spacer(doc)
checkbox(doc, "PRESTA IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i propri dati personali per le finalità e "
              "secondo le modalità ivi descritte")
spacer(doc)
checkbox(doc, "NEGA IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i propri dati personali per le finalità e "
              "secondo le modalità ivi descritte")
spacer(doc)
checkbox(doc, "AUTORIZZA l'Università degli Studi di Genova all'utilizzo "
              "delle proprie immagini e/o della propria voce per "
              "realizzazione di video e materiali multimediali, "
              "realizzati e utilizzati esclusivamente per finalità di "
              "ricerca scientifica (da compilare in caso di presenza di "
              "registrazioni audio e/o video)")

spacer(doc, 2)
body(doc, "Data ___/___/______")
spacer(doc)
body(doc, "Nome _________________________     Cognome _________________________")
spacer(doc)
body(doc, "Firma leggibile ________________________________________")

doc.add_page_break()

# ===================== SEZIONE B — MINORI IN GRADO DI COMPRENDERE =====================
hcenter(doc, "UNIVERSITÀ DEGLI STUDI DI GENOVA", bold=True, size=14)
hcenter(doc, "Dipartimento di Neuroscienze, Riabilitazione, Oftalmologia, "
             "Genetica e Scienze Materno-Infantili")
spacer(doc)

hcenter(doc,
    "CONSENSO INFORMATO PER IL TRATTAMENTO DEI DATI PERSONALI "
    "NELL'AMBITO DEI PROGETTI DI RICERCA",
    bold=True, size=11)
hcenter(doc,
    "da compilare nel caso di raccolta di dati personali di minori in "
    "grado di comprendere l'informativa",
    bold=True, size=10)
spacer(doc)

body(doc,
    "Nota: lo studio ARISE NON coinvolge soggetti minorenni. La presente "
    "sezione è inclusa per completezza del template e non è applicabile "
    "al presente studio.", italic=True)
spacer(doc)

body(doc, "Il/la sottoscritto/a ………………………………………………………………………")
spacer(doc)

checkbox(doc, "DICHIARA di aver preso visione dell'Informativa per il "
              "trattamento dei dati personali di cui alla sezione A")
spacer(doc)
body(doc, "Data ___/___/______        Firma leggibile ___________________________")

spacer(doc)
body(doc, "In qualità di rappresentante legale")
checkbox(doc, "PRESTANO IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i dati personali per le finalità e secondo le "
              "modalità ivi descritte")
spacer(doc)
checkbox(doc, "NEGANO IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i dati personali per le finalità e secondo le "
              "modalità ivi descritte")
spacer(doc)
checkbox(doc, "AUTORIZZANO l'Università degli Studi di Genova all'utilizzo "
              "delle proprie immagini e/o della propria voce per "
              "realizzazione di video e materiali multimediali, realizzati "
              "e utilizzati esclusivamente per finalità di ricerca "
              "scientifica (da compilare in caso di presenza di "
              "registrazioni audio e/o video)")

spacer(doc, 2)
body(doc, "Data ___/___/______")
spacer(doc)
body(doc, "Nome ____________________   Cognome ____________________")
spacer(doc)
body(doc, "Firma leggibile ________________________________________")
spacer(doc)
body(doc, "Nome ____________________   Cognome ____________________")
spacer(doc)
body(doc, "Firma leggibile ________________________________________")

doc.add_page_break()

# ===================== SEZIONE B — MINORI =====================
hcenter(doc, "UNIVERSITÀ DEGLI STUDI DI GENOVA", bold=True, size=14)
hcenter(doc, "Dipartimento di Neuroscienze, Riabilitazione, Oftalmologia, "
             "Genetica e Scienze Materno-Infantili")
spacer(doc)

hcenter(doc,
    "CONSENSO INFORMATO PER IL TRATTAMENTO DEI DATI PERSONALI "
    "NELL'AMBITO DEI PROGETTI DI RICERCA",
    bold=True, size=11)
hcenter(doc,
    "da compilare nel caso di raccolta di dati personali di minori",
    bold=True, size=10)
spacer(doc)

body(doc,
    "Nota: lo studio ARISE NON coinvolge soggetti minorenni. La presente "
    "sezione è inclusa per completezza del template e non è applicabile "
    "al presente studio.", italic=True)
spacer(doc)

body(doc,
    "I sottoscritti, in qualità di rappresentanti legali di "
    "………………………………………………………………………")
spacer(doc)

checkbox(doc, "DICHIARANO di aver preso visione dell'Informativa per il "
              "trattamento dei dati personali di cui alla sezione A)")
spacer(doc)
checkbox(doc, "PRESTANO IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i dati personali del proprio figlio per le "
              "finalità e secondo le modalità ivi descritte")
spacer(doc)
checkbox(doc, "NEGANO IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i dati personali del proprio figlio per le "
              "finalità e secondo le modalità ivi descritte")
spacer(doc)
checkbox(doc, "AUTORIZZANO l'Università degli Studi di Genova all'utilizzo "
              "delle immagini e/o della voce del proprio figlio per la "
              "realizzazione di video e materiali multimediali, realizzati "
              "e utilizzati esclusivamente per finalità di ricerca "
              "scientifica (da compilare in caso di presenza di "
              "registrazioni audio e/o video)")

spacer(doc, 2)
body(doc, "Data ___/___/______")
spacer(doc)
body(doc, "Nome ____________________   Cognome ____________________")
spacer(doc)
body(doc, "Firma leggibile ________________________________________")
spacer(doc)
body(doc, "Nome ____________________   Cognome ____________________")
spacer(doc)
body(doc, "Firma leggibile ________________________________________")

doc.save(OUT)
print(f"Wrote {OUT}")
