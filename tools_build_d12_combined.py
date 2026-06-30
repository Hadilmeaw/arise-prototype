"""Build the combined ARISE D1.2 Ethics Committee submission package as a single .docx.

The deliverable contains four documents in sequence, followed by a
single consolidated Annexes section at the end:
  01 Clinical Investigation Plan          (body)
  02 Investigator's Brochure              (body)
  03 Patient Information Sheet            (body)
  04 Informed Consent Form                (body)
  Annexes (all parts):
    A.1 Schedule of assessments           (from 01)
    A.2 Glossary                          (from 01)
    B.1 Applicable standards              (from 02)
    D.1 Modulo di Consenso Informato (italiano) (from 04)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY = RGBColor(0x1A, 0x33, 0x5C)
INK  = RGBColor(0x14, 0x14, 0x14)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
LIGHT = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


LINK_RE = re.compile(r"\[(.+?)\]\((.+?)\)")
IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.+?)\)$")


def add_runs(p, text):
    pattern = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = p.add_run(token[2:-2]); r.bold = True
        elif token.startswith("`"):
            r = p.add_run(token[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        elif token.startswith("["):
            lm = LINK_RE.match(token)
            r = p.add_run(lm.group(1))
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            r.underline = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_md_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    has_header = any(c.strip() for c in rows[0])
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    if not has_header:
        rows = rows[1:]
        tbl._element.remove(tbl.rows[0]._tr)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, val.strip())
            if has_header and i == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "1A335C")
                for run in p.runs:
                    run.font.color.rgb = WHITE
            elif not has_header and j == 0:
                for run in p.runs:
                    run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def parse_table_block(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line):
            i += 1; continue
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "": cells = cells[1:]
        if cells and cells[-1] == "": cells = cells[:-1]
        rows.append(cells)
        i += 1
    return rows, i


def render_markdown(doc, md_text, root_dir, *, suppress_first_h1=True,
                    suppress_first_front_matter=True):
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    skipped_title = not suppress_first_h1
    skipped_front_matter_table = not suppress_first_front_matter

    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1; continue
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            i += 1; continue
        im = IMAGE_RE.match(stripped)
        if im:
            alt, src = im.group(1), im.group(2)
            img_path = (root_dir / src) if not Path(src).is_absolute() else Path(src)
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                try:
                    run.add_picture(str(img_path), width=Inches(5.5))
                except Exception:
                    p.add_run(f"[Image: {alt}]").italic = True
            i += 1; continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            ttl = m.group(2).strip()
            ttl = re.sub(r"^\*\*(.+)\*\*$", r"\1", ttl)
            if level == 1 and not skipped_title:
                skipped_title = True
                i += 1; continue
            p = doc.add_heading(level=min(level, 4))
            add_runs(p, ttl)
            i += 1; continue
        if stripped.startswith("|"):
            rows, j = parse_table_block(lines, i)
            if not skipped_front_matter_table and len(rows) >= 3 \
               and rows[0] and len(rows[0]) == 2 and rows[0][0].strip().lower() in ("", "field"):
                skipped_front_matter_table = True
                i = j; continue
            add_md_table(doc, rows)
            i = j; continue
        if re.match(r"^[-*+]\s+", stripped):
            while i < n and re.match(r"^[-*+]\s+", lines[i].strip()):
                item = re.sub(r"^[-*+]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, item)
                i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            counter = 0
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                counter += 1
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                r = p.add_run(f"{counter}. ")
                r.bold = True
                add_runs(p, item)
                i += 1
            continue
        if stripped.startswith("> "):
            content = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            r = p.add_run(content)
            r.italic = True
            r.font.color.rgb = GREY
            i += 1; continue
        if stripped.startswith("```"):
            i += 1; code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code))
            r.font.name = "Consolas"; r.font.size = Pt(9)
            continue
        para_lines = [line]; i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#|\||[-*+]\s|\d+\.\s|>\s|```|-{3,}|!\[)", lines[i].strip()
        ):
            para_lines.append(lines[i]); i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        add_runs(p, para_text)


# Split a markdown into (body_lines, [(annex_title, annex_lines), ...])

def split_body_and_annexes(md_text):
    lines = md_text.splitlines()
    annex_start = None
    for i, line in enumerate(lines):
        if re.match(r"^##\s+Annex\s+", line.strip()):
            annex_start = i
            break
    if annex_start is None:
        return "\n".join(lines), []
    body = "\n".join(lines[:annex_start])
    annexes = []
    cur_title = None
    cur_lines = []
    for line in lines[annex_start:]:
        m = re.match(r"^##\s+Annex\s+\S+?:\s*(.*)$", line.strip())
        if m:
            if cur_title is not None:
                annexes.append((cur_title, "\n".join(cur_lines).strip()))
            cur_title = m.group(1).strip()
            cur_lines = []
        else:
            cur_lines.append(line)
    if cur_title is not None:
        annexes.append((cur_title, "\n".join(cur_lines).strip()))
    return body, annexes


# Chrome

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def set_page_footer(section, text_left, page_label):
    footer = section.footer
    for p in list(footer.paragraphs):
        if p.text:
            for r in p.runs: r.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = fp.add_run(text_left)
    r1.font.size = Pt(9); r1.font.color.rgb = GREY
    r2 = fp.add_run("\t\t")
    r3 = fp.add_run(page_label + "  ")
    r3.font.size = Pt(9); r3.font.color.rgb = GREY
    run = fp.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.size = Pt(9); run.font.color.rgb = GREY


def add_text(doc, text, *, size=11, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color
    return p


def build_cover_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    add_text(doc, "ARISE", size=48, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_text(doc,
             "Augmented Rehabilitation & Intelligent System for Enhancement",
             size=14, color=GREY, italic=True, space_after=20)
    add_text(doc, "D1.2 Ethics Committee Submission Package", size=28, bold=True, color=INK,
             space_after=4)
    add_text(doc, "Clinical investigation of the ARISE system. Class IIa under MDR Annex VIII Rule 11",
             size=14, color=GREY, italic=True, space_after=40)

    add_text(doc, "Documents in this submission package", size=12, bold=True, color=NAVY,
             space_after=8)
    docs = [
        ("01", "Clinical Investigation Plan (CIP)",
         "The protocol. Objectives, design, endpoints, population, procedures, statistical considerations, adverse-event reporting under MDR Article 80"),
        ("02", "Investigator's Brochure (IB)",
         "Device description, principle of operation, risk analysis, instructions for the investigator"),
        ("03", "Patient Information Sheet (PIS)",
         "Plain-language information for participants. Explains the investigation, what participation involves, voluntary nature, data protection"),
        ("04", "Informed Consent Form (ICF)",
         "Tick-box consent with statements, specific consents, signatures. Italian version included in the Annexes at the end"),
        ("•", "Annexes (all parts)",
         "Schedule of assessments, glossary, applicable standards, Italian Modulo di Consenso Informato"),
    ]
    for num, title, sub in docs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r1 = p.add_run(num + "   ")
        r1.font.name = "Calibri"; r1.font.size = Pt(14); r1.font.bold = True
        r1.font.color.rgb = NAVY
        r2 = p.add_run(title)
        r2.font.name = "Calibri"; r2.font.size = Pt(14); r2.font.bold = True
        r2.font.color.rgb = INK
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.9)
        p2.paragraph_format.space_after = Pt(10)
        r3 = p2.add_run(sub)
        r3.font.name = "Calibri"; r3.font.size = Pt(11); r3.font.color.rgb = GREY

    for _ in range(2):
        doc.add_paragraph()
    add_text(doc, "Project context", size=10, bold=True, color=NAVY)
    meta = [
        ("Funding instrument", "ERDF"),
        ("Deliverable", "D1.2"),
        ("Work Package", "WP1, Task T1.3"),
        ("Planned delivery", "Month 3 (mid-September 2026)"),
        ("Manufacturer and sponsor", "Innovina S.r.l."),
        ("Scientific partner", "DINOGMI, University of Genoa"),
        ("Clinical site", "Studio Buccarella"),
        ("Reference instrumentation site", "REHELAB, University of Genoa"),
        ("Confirmed MDR class", "Class IIa under Annex VIII Rule 11"),
        ("Phase 1 regulatory framework", "MDR Articles 62 to 82, Annex XV"),
        ("Languages", "English master, Italian for the participant-facing forms"),
        ("Deliverable version", "2.1, 9 June 2026"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Light List Accent 1"
    for i, (k, v) in enumerate(meta):
        c1 = tbl.rows[i].cells[0]
        c1.text = ""
        p = c1.paragraphs[0]
        r = p.add_run(k)
        r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = INK
        c2 = tbl.rows[i].cells[1]
        c2.text = ""
        p = c2.paragraphs[0]
        r = p.add_run(v)
        r.font.name = "Calibri"; r.font.size = Pt(10)
        r.font.color.rgb = INK


def build_master_toc(doc):
    add_page_break(doc)
    add_text(doc, "Contents", size=24, bold=True, color=NAVY, space_after=20)

    items = [
        ("01", "Clinical Investigation Plan"),
        ("02", "Investigator's Brochure"),
        ("03", "Patient Information Sheet"),
        ("04", "Informed Consent Form"),
        ("•", "Annexes (consolidated): A.1, A.2 from Part 01; B.1 from Part 02; D.1 from Part 04"),
    ]
    for num, title in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r1 = p.add_run(num + "    ")
        r1.font.name = "Calibri"; r1.font.size = Pt(14); r1.font.bold = True
        r1.font.color.rgb = NAVY
        r2 = p.add_run(title)
        r2.font.name = "Calibri"; r2.font.size = Pt(14); r2.font.color.rgb = INK

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Companion documents (not included in this file): D1.1 (clinical requirements and KPIs), "
        "the ARISE Compliance Dossier (DMP + DPIA + MDR Compliance Plan), "
        "and the Validation Acceptance Criteria (primary endpoint threshold)."
    )
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.italic = True
    r.font.color.rgb = GREY


def build_hero_page(doc, number, title, subtitle, summary_lines):
    add_page_break(doc)
    for _ in range(3):
        doc.add_paragraph()
    add_text(doc, number, size=72, bold=True, color=NAVY, space_after=4)
    add_text(doc, title, size=28, bold=True, color=INK, space_after=4)
    add_text(doc, subtitle, size=14, color=GREY, italic=True, space_after=24)
    add_text(doc, "Document summary", size=11, bold=True, color=NAVY,
             space_after=6)
    for ln in summary_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("•  " + ln)
        r.font.name = "Calibri"; r.font.size = Pt(11); r.font.color.rgb = INK
    add_page_break(doc)


def render_annex(doc, annex_label, annex_title, annex_md, root):
    """Render a single annex with relabeled heading."""
    # Heading: Annex {label}: {title}
    add_page_break(doc)
    p = doc.add_heading(level=1)
    add_runs(p, f"Annex {annex_label}: {annex_title}")
    # Render body content using the markdown renderer
    render_markdown(doc, annex_md, root,
                    suppress_first_h1=False, suppress_first_front_matter=False)


# Main build

def build(out_path: Path):
    root = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2")
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for lvl, size in [(1, 18), (2, 14), (3, 12), (4, 11)]:
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = NAVY
        s.font.bold = True

    set_page_footer(section, "ARISE D1.2 Submission Package", "Page")

    build_cover_page(doc)
    build_master_toc(doc)

    # Load and split each document
    cip_md = (root / "ARISE_D1.2_Clinical_Investigation_Plan.md").read_text(encoding="utf-8")
    ib_md  = (root / "ARISE_D1.2_Investigators_Brochure.md").read_text(encoding="utf-8")
    pis_md = (root / "ARISE_D1.2_Patient_Information_Sheet.md").read_text(encoding="utf-8")
    icf_md = (root / "ARISE_D1.2_Informed_Consent_Form.md").read_text(encoding="utf-8")

    cip_body, cip_annexes = split_body_and_annexes(cip_md)
    ib_body,  ib_annexes  = split_body_and_annexes(ib_md)
    pis_body, pis_annexes = split_body_and_annexes(pis_md)
    icf_body, icf_annexes = split_body_and_annexes(icf_md)

    # 01 - CIP
    build_hero_page(
        doc,
        number="01",
        title="Clinical Investigation Plan",
        subtitle="The protocol submitted to the Ethics Committee",
        summary_lines=[
            "ARISE classified as Class IIa under MDR Annex VIII Rule 11.",
            "Phase 1 first-in-clinic demonstration. Single-site, single-arm, within-subject reference comparison.",
            "Two-stage data collection: corpus building (camera + DINOGMI annotation), then live ARISE deployment with biofeedback.",
            "Primary endpoint expressed as agreement between an ARISE-derived KPI and the simultaneous reference at REHELAB. Exact threshold deferred to the Validation Acceptance Criteria.",
            "Adverse-event reporting under MDR Article 80: 2 calendar days for imminent risk, 7 calendar days otherwise.",
        ],
    )
    render_markdown(doc, cip_body, root)

    # 02 - IB
    build_hero_page(
        doc,
        number="02",
        title="Investigator's Brochure",
        subtitle="Device description and risk summary",
        summary_lines=[
            "Markerless single-camera system, supervised by a qualified rehabilitation professional.",
            "Three-layer detection: population rule, per-patient z-score, deep learning model.",
            "Identified hazards table with mitigation and residual risk per hazard.",
            "Anticipated adverse events limited to events from the Sit-to-Stand transfer itself, not from the device.",
            "Instructions for the investigator: training, set-up, session use, end-of-session, malfunction handling.",
        ],
    )
    render_markdown(doc, ib_body, root)

    # 03 - PIS
    build_hero_page(
        doc,
        number="03",
        title="Patient Information Sheet",
        subtitle="Plain-language information for prospective participants",
        summary_lines=[
            "Explains the purpose of the investigation, what participation involves, and the voluntary nature of participation.",
            "Sessions: typically 4 to 8, each lasting 30 to 45 minutes, over 2 to 4 weeks.",
            "How personal data is collected, who has access, where it is stored, how long it is kept.",
            "Participant rights and how to exercise them.",
            "Contact information and insurance.",
        ],
    )
    render_markdown(doc, pis_body, root)

    # 04 - ICF
    build_hero_page(
        doc,
        number="04",
        title="Informed Consent Form",
        subtitle="English master, with the Italian version in the Annexes",
        summary_lines=[
            "Ten statements (S1 to S10) initialled by the participant.",
            "Six specific consents (C1 to C6), independent. C1, C2, C6 necessary for participation.",
            "Signatures of participant, legal representative (where applicable), and Principal Investigator.",
            "Withdrawal-of-consent block at the end, with choice on the disposition of already-collected data.",
            "Italian translation appears in the consolidated Annexes section at the end of this file.",
        ],
    )
    render_markdown(doc, icf_body, root)

    # Annexes consolidated section
    build_hero_page(
        doc,
        number="❖",
        title="Annexes",
        subtitle="Consolidated annexes for all parts of the submission package",
        summary_lines=[
            "A.1, A.2  -  from Part 01 Clinical Investigation Plan",
            "B.1       -  from Part 02 Investigator's Brochure",
            "D.1       -  from Part 04 Informed Consent Form (Italian translation, signed by the participant)",
        ],
    )

    # Render each annex with new labels
    # CIP annexes -> A.1, A.2
    for idx, (title, content) in enumerate(cip_annexes, start=1):
        render_annex(doc, f"A.{idx}", title, content, root)
    # IB annexes -> B.1, ...
    for idx, (title, content) in enumerate(ib_annexes, start=1):
        render_annex(doc, f"B.{idx}", title, content, root)
    # PIS annexes -> C.1, ... (currently none)
    for idx, (title, content) in enumerate(pis_annexes, start=1):
        render_annex(doc, f"C.{idx}", title, content, root)
    # ICF annexes -> D.1, ...
    for idx, (title, content) in enumerate(icf_annexes, start=1):
        render_annex(doc, f"D.{idx}", title, content, root)

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    out = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2\ARISE_D1.2_Submission_Package.docx")
    build(out)
