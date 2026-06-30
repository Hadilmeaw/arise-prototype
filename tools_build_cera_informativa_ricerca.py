"""ARISE — Informativa Ricerca (table-style format), matching the
modern UNIGE template "Informativaricerca.docx" with table cells for
TITOLARE / DPO / CONTATTI / BASE GIURIDICA / CATEGORIA / OBBLIGO /
MODALITA / PROCESSO AUTOMATIZZATO / TEMPO / DESTINATARI / TRASFERIMENTO
ESTERO / DIRITTI / RECLAMO / ULTERIORI INFORMAZIONI. Followed by Sezione B
with consent checkboxes for adults, minors-capable, and minors.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2")
OUT  = ROOT / "ARISE_CERA_Informativa_Ricerca.docx"

STUDY_TITLE = (
    "ARISE — A feasibility study on markerless video-based biomechanical "
    "assessment of the Sit-to-Stand task in healthy adults and older adults"
)

def hcenter(doc, t, *, bold=False, size=11, italic=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size)

def section(doc, t, *, size=12):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = True; r.font.size = Pt(size)

def body(doc, t, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(t); r.italic = italic; r.font.size = Pt(size)

def checkbox(doc, t):
    p = doc.add_paragraph()
    r = p.add_run("☐  " + t); r.font.size = Pt(11)

def spacer(doc, n=1):
    for _ in range(n): doc.add_paragraph()

def set_cell(cell, text, *, bold=False, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)

def add_kv_row(table, key, value, *, key_size=11, val_size=10.5):
    row = table.add_row().cells
    set_cell(row[0], key, bold=True, size=key_size)
    set_cell(row[1], value, size=val_size)

# ===========================================================================
doc = Document()
s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(2.0)

# ===================== HEADER =====================
hcenter(doc,
    "INFORMATIVA PER IL TRATTAMENTO DI DATI PERSONALI "
    "NELL'AMBITO DEI PROGETTI DI RICERCA", bold=True, size=12)
hcenter(doc, STUDY_TITLE, bold=True, italic=True, size=11)
hcenter(doc,
    "resa ai sensi dell'art. 13 del Regolamento UE n. 2016/679 (GDPR)",
    italic=True, size=10)
spacer(doc)

# ===================== SEZIONE A =====================
section(doc, "Sezione A")
spacer(doc)

body(doc, "Gentile Interessato/a,")
body(doc,
    "Conformemente alla normativa vigente, l'Università degli Studi di "
    "Genova, nel rispetto della tutela della Sua riservatezza e dei Suoi "
    "diritti, impronta il trattamento dei dati personali che La "
    "riguardano ai principi di cui all'art. 5 del GDPR, tra i quali "
    "liceità, correttezza e trasparenza, adeguatezza, pertinenza e "
    "limitazione, esattezza e aggiornamento, non eccedenza e "
    "responsabilizzazione.")
body(doc, "Ai sensi dell'art. 13 GDPR, Le forniamo le seguenti informazioni:")
spacer(doc)

# ===================== INFORMATIVA TABLE =====================
table = doc.add_table(rows=0, cols=2)
table.style = "Light Grid Accent 1"
# Column widths
for r in table.rows:
    r.cells[0].width = Cm(5.0)
    r.cells[1].width = Cm(12.0)

add_kv_row(table, "TITOLARE DEL TRATTAMENTO DEI DATI",
    "Università degli studi di Genova, nella persona del Rettore "
    "pro tempore.\n"
    "I dati di contatto sono:\n"
    "Rettorato – Genova, Via Balbi, n. 5\n"
    "Telefono: (+39) 010209-9221, (+39) 010209-51929\n"
    "e-mail: rettore@unige.it    PEC: protocollo@pec.unige.it")

add_kv_row(table, "DATA PROTECTION OFFICER (DPO)",
    "Presso il titolare del trattamento è presente il responsabile "
    "della protezione dei dati (DPO).\n"
    "I dati di contatto sono:\n"
    "Telefono: (+39) 3385021237\n"
    "e-mail: dpo@unige.it\n"
    "PEC: protocollo@pec.liguriadigitale.it")

add_kv_row(table, "CONTATTI E SEGNALAZIONI",
    "Per informazioni circa il progetto o eventuali segnalazioni è "
    "possibile contattare:\n"
    "•  Il Responsabile del progetto: Prof. Marco Testa "
    "(email: marco.testa@unige.it; telefono: +39 331 2611548)\n"
    "•  Il Coordinatore del progetto: Prof. Marco Testa "
    "(email: marco.testa@unige.it)\n"
    "•  Altro contatto: Dott. Mirko Job "
    "(email: mirko.job.1991@gmail.com; telefono: +39 349 5199778)")

add_kv_row(table, "BASE GIURIDICA E FINALITÀ DEL TRATTAMENTO",
    "Il trattamento dei dati personali, nei limiti delle finalità "
    "perseguite, è effettuato in quanto necessario per l'esecuzione di "
    "un compito di interesse pubblico o connesso all'esercizio di "
    "pubblici poteri di cui è investito il titolare del trattamento "
    "(art. 6, par. 1, lett. e) del GDPR).\n"
    "Il trattamento dei dati particolari avviene inoltre sulla base del "
    "consenso esplicito al trattamento di tali dati (art. 9, par. 2, "
    "lett. a) del GDPR) e delle finalità istituzionali dell'Ateneo "
    "nell'ambito della ricerca scientifica (art. 9, par. 2, lett. g) "
    "del GDPR e art. 2-sexies, co. 2, lett. cc) del D. Lgs. n. 196/2003 "
    "\"codice privacy\").\n"
    "Lei ha il diritto di revocare il consenso in qualsiasi momento. La "
    "revoca del consenso non pregiudica la liceità del trattamento "
    "basata sul consenso prestato prima della revoca. Non saranno "
    "raccolti ulteriori dati che La riguardano, ferma restando "
    "l'utilizzazione di quelli eventualmente già raccolti per "
    "determinare, senza alterarli, i risultati della ricerca.\n"
    f"I dati da Lei forniti verranno trattati per finalità di ricerca "
    f"scientifica, nell'ambito del progetto di ricerca dal titolo "
    f"\"{STUDY_TITLE}\".\n"
    "La ricerca è finalizzata a (1) raccogliere dati biomeccanici sul "
    "Sit-to-Stand in una popolazione di adulti (>18 anni), utilizzando "
    "registrazioni video da telecamere RGB e/o RGB-D analizzate "
    "attraverso metodologie non invasive di pose estimation senza "
    "marker; (2) caratterizzare i principali indicatori cinematici "
    "(KPI) del movimento; (3) validare i dati ottenuti attraverso "
    "confronto con misure di riferimento ottenute attraverso sistemi "
    "di analisi cinematica e cinetica del movimento gold-standard "
    "validati con metodologie non invasive.\n"
    "I dati saranno raccolti presso i Laboratori di ricerca del "
    "Dipartimento DINOGMI (Campus Universitario di Savona) durante 2 "
    "sessioni di acquisizione della durata di 10-15 minuti ciascuna, "
    "in cui il/la partecipante eseguirà alcune ripetizioni del "
    "Sit-to-Stand secondo protocolli standardizzati (5xSTS e 30-second "
    "Chair Stand Test), registrate mediante telecamere RGB e/o RGB-D.")

add_kv_row(table, "CATEGORIA DI DATI TRATTATI",
    "I dati che saranno oggetto del trattamento nell'ambito del "
    "presente studio appartengono alle seguenti categorie:\n"
    "•  dati personali, quali nome e cognome, data di nascita, indirizzo "
    "e-mail personale, numero di telefono/cellulare;\n"
    "•  dati antropometrici e sociodemografici (età, altezza, peso, "
    "sesso assegnato alla nascita, genere);\n"
    "•  immagini e voce per videoregistrazioni acquisite durante "
    "l'esecuzione del task Sit-to-Stand;\n"
    "•  dati cinematici biomeccanici derivati dall'analisi video "
    "(keypoint anatomici, angoli articolari, velocità, durate di fase);\n"
    "•  categorie particolari di dati personali di cui all'art. 9, "
    "paragrafo 1, del GDPR, quali dati anamnestici idonei a rivelare lo "
    "stato di salute generale del/della partecipante.")

add_kv_row(table, "OBBLIGO DI CONFERIMENTO DEI DATI",
    "Il conferimento dei dati è facoltativo, cioè non discende da un "
    "obbligo normativo, e l'eventuale rifiuto di fornire tali dati non "
    "ha conseguenze. Il conferimento dei dati per le finalità sopra "
    "indicate è indispensabile allo svolgimento dello studio. Il "
    "rifiuto di conferirli non consentirà all'Interessato di "
    "partecipare allo studio in parola.\n"
    "Il conferimento dei dati di contatto è facoltativo, cioè non "
    "discende da un obbligo normativo, ma è necessario per consentire "
    "la conservazione dei dati per un periodo più lungo di quello "
    "previsto per la conclusione del presente studio ed eventualmente "
    "permettere al Titolare di ricontattare l'Interessato affinché "
    "possa esprimere, se lo riterrà, un nuovo specifico consenso per "
    "una ulteriore ricerca. Il mancato conferimento dei dati per tali "
    "finalità avrà come unica conseguenza l'impossibilità di "
    "realizzare quanto da ultimo descritto.")

add_kv_row(table, "MODALITÀ DI TRATTAMENTO",
    "Le attività di trattamento dei dati personali saranno effettuate "
    "con modalità cartacea e informatizzata, rispettando i principi di "
    "necessità, liceità, correttezza, esattezza, proporzionalità, "
    "minimizzazione, pertinenza e non eccedenza, anche da parte dei "
    "soggetti appositamente autorizzati al trattamento ai sensi degli "
    "artt. 28 e 29 del GDPR, e adottando le adeguate misure tecniche e "
    "organizzative previste dall'art. 32 dello stesso Regolamento.\n"
    "Il Responsabile che conduce lo studio e i collaboratori da questi "
    "nominati provvederanno a identificare il/la partecipante con un "
    "codice univoco randomico (es. sbj4112), generato automaticamente "
    "dal software di acquisizione fornito da Innovina S.r.l., al fine "
    "di consentire che la Sua identità non sia conosciuta dagli altri "
    "soggetti coinvolti nello studio.")

add_kv_row(table, "PROCESSO AUTOMATIZZATO",
    "Il titolare non utilizza processi decisionali automatizzati "
    "finalizzati alla profilazione individuale degli interessati. I "
    "dati cinematici acquisiti sono soggetti a elaborazione automatica "
    "(pose estimation, estrazione di indicatori biomeccanici, "
    "addestramento di modelli di machine learning a fini di ricerca), "
    "ma tale elaborazione è esclusivamente di natura statistico-"
    "scientifica e non produce decisioni con effetti giuridici o "
    "incidenti significativamente sull'interessato.")

add_kv_row(table, "TEMPO DI CONSERVAZIONE DEI DATI PERSONALI",
    "I dati personali, raccolti nell'ambito del progetto, saranno "
    "conservati nelle forme sopra indicate per 3 anni dopo il termine "
    "dello studio.\n"
    "I dati raccolti saranno comunque conservati per i tempi stabiliti "
    "dalla normativa vigente o dai regolamenti d'Ateneo e comunque in "
    "una forma che consenta l'identificazione degli interessati per un "
    "arco di tempo non superiore al conseguimento delle finalità per "
    "le quali sono trattati. I dati identificativi diretti (nome, "
    "cognome, telefono, e-mail) sono conservati esclusivamente presso "
    "il Coordinatore del progetto, in archivio cifrato, all'interno "
    "del Dipartimento DINOGMI. Le registrazioni video grezze e i dati "
    "cinematici sono archiviati sul server protetto del Dipartimento "
    "DINOGMI e conservati per il tempo strettamente necessario "
    "all'analisi biomeccanica (massimo 12 mesi) e successivamente "
    "cancellate, salvo consenso esplicito alla conservazione per "
    "ricerca futura.")

add_kv_row(table, "DESTINATARI DEI DATI PERSONALI",
    "I dati personali saranno trattati all'interno dell'Ateneo da "
    "soggetti autorizzati dal titolare debitamente formati e istruiti, "
    "nel rispetto del segreto d'ufficio o da Responsabili ex art. 28 "
    "GDPR. In particolare, le persone autorizzate al trattamento sono:\n"
    "•  Prof. Marco Testa, Coordinatore del progetto, accesso ai dati "
    "identificativi diretti e ai dati pseudonimizzati;\n"
    "•  Dottorando del gruppo di ricerca, accesso ai dati "
    "pseudonimizzati.\n"
    "Innovina S.r.l., partner tecnico-industriale del progetto ARISE, "
    "è designata Responsabile esterno del trattamento ai sensi "
    "dell'Art. 28 GDPR e vincolata da apposito Data Processing "
    "Agreement (DPA) attualmente in corso di formalizzazione: potrà "
    "ricevere esclusivamente i dati delle acquisizioni in forma "
    "pseudonimizzata, ai fini dell'elaborazione tecnica e "
    "dell'addestramento dei modelli previsti dal progetto. Nessun "
    "dato identificativo del/della partecipante sarà trasmesso a "
    "Innovina S.r.l.\n"
    "I dati personali saranno comunque oggetto di comunicazione nei "
    "confronti di soggetti, enti o Autorità verso i quali la "
    "comunicazione sia obbligatoria in forza di disposizioni di legge "
    "o di regolamento.\n"
    "La diffusione dei dati scientifici risultanti dalle analisi dei "
    "dati dello Studio potrà avvenire solo in forma anonima, aggregata "
    "e per sole finalità scientifiche dello Studio. I risultati delle "
    "ricerche scientifiche potranno essere presentati in forma "
    "aggregata nell'ambito di Convegni o pubblicati su riviste "
    "specializzate senza mai permettere la precisa identificazione "
    "dell'interessato.")

add_kv_row(table, "TRASFERIMENTO DI DATI ALL'ESTERO",
    "I dati personali non saranno trasferiti in paesi ubicati al di "
    "fuori dell'Unione europea. Tutti i dati restano all'interno dello "
    "Spazio Economico Europeo, conservati presso le sedi del "
    "Dipartimento DINOGMI (Campus Universitario di Savona) e presso "
    "l'infrastruttura informatica fornita da Innovina S.r.l. in "
    "regione europea.")

add_kv_row(table, "DIRITTI DELL'INTERESSATO",
    "Sono diritti dell'interessato:\n"
    "•  l'accesso ai dati (art. 15 GDPR);\n"
    "•  la rettifica dei dati (art. 16 GDPR);\n"
    "•  la cancellazione dei dati (c.d. \"diritto all'oblio\"), salvo "
    "che per i dati contenuti negli atti che devono obbligatoriamente "
    "essere conservati dall'Università (art. 17 GDPR);\n"
    "•  la limitazione di trattamento (art. 18 GDPR);\n"
    "•  la portabilità dei dati (art. 20 GDPR);\n"
    "•  l'opposizione al trattamento (art. 21 GDPR).\n"
    "L'art. 21 del GDPR prevede il diritto di opporsi in qualsiasi "
    "momento, per motivi connessi alla situazione particolare "
    "dell'interessato, al trattamento dei dati personali che lo "
    "riguardano. L'opposizione non pregiudica la liceità del "
    "trattamento svolto precedentemente.")

add_kv_row(table, "RECLAMO",
    "Fatto salvo ogni altro ricorso amministrativo e giurisdizionale, "
    "qualora ritenga che il trattamento dei dati non sia conforme a "
    "quanto previsto dal GDPR, ha diritto, ai sensi dell'art. 77 GDPR, "
    "di avanzare un reclamo al Garante italiano per la protezione dei "
    "dati personali (www.garanteprivacy.it) o all'Autorità garante "
    "dello Stato UE in cui risiede abitualmente o lavora oppure del "
    "luogo ove si è verificata la presunta violazione.")

add_kv_row(table, "ULTERIORI INFORMAZIONI",
    "L'informativa è messa a disposizione dell'interessato, unitamente "
    "ai suoi eventuali aggiornamenti.\n"
    "Ulteriori informazioni riguardanti il trattamento dei dati "
    "personali, incluse le modalità per l'esercizio dei diritti, sono "
    "reperibili sul sito web https://unige.it/privacy.\n"
    "Ultimo aggiornamento all'informativa: [MESE ANNO da inserire].")

spacer(doc)
body(doc,
    "Seguono istruzioni applicabili nella maggioranza dei casi. In caso "
    "di dubbi rivolgersi all'ufficio privacy anche in corso di redazione "
    "dell'informativa.", italic=True)
spacer(doc)

# ===================== PRESA VISIONE =====================
hcenter(doc, "PRESA VISIONE", bold=True, size=11)
spacer(doc)
body(doc,
    "Laddove il trattamento si fondi sulle finalità istituzionali, o su "
    "altre finalità diverse dal consenso, è opportuno prevedere una "
    "spunta di presa visione dell'informativa, secondo le modalità che "
    "seguono:")

p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
r = p.add_run("•  "); r.font.size = Pt(11)
r = p.add_run("ove l'informativa sia fornita digitalmente "); r.bold = True; r.font.size = Pt(11)
r = p.add_run(
    "(es. prima dell'accesso a un questionario su Microsoft Form, prima "
    "dell'iscrizione a qualche servizio, etc.), sarà sufficiente una "
    "check box dal tenore \"Dichiaro di aver preso visione "
    "dell'informativa [nome informativa, inserendo il link "
    "all'informativa fornito dall'ufficio privacy al termine della "
    "redazione della stessa]\". La check box dovrà essere bloccante, "
    "ossia non dovrà essere possibile proseguire senza averla sbarrata;")
r.font.size = Pt(11)

p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
r = p.add_run("•  "); r.font.size = Pt(11)
r = p.add_run("ove l'informativa sia fornita cartacea, "); r.bold = True; r.font.size = Pt(11)
r = p.add_run(
    "l'indicazione \"Dichiaro di aver preso visione dell'informativa "
    "[nome informativa]\" deve essere inserita in calce a una copia "
    "dell'informativa che resterà all'Ateneo, lasciando un'altra copia "
    "nella disponibilità dell'interessato.")
r.font.size = Pt(11)

spacer(doc)

# ===================== CONSENSO =====================
hcenter(doc, "CONSENSO", bold=True, size=11)
spacer(doc)
body(doc,
    "Nel caso in cui è richiesto il consenso (ossia necessariamente in "
    "presenza di dati particolari e di dati di minori, ma anche in "
    "tutti quei casi in cui la base giuridica è data dall'art. 6, par. "
    "1, lett. a), oltre alla presa visione dell'informativa, è "
    "necessario acquisire il consenso. Ciò potrà avvenire:")

# Digital — adults
p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
r = p.add_run("•  "); r.font.size = Pt(11)
r = p.add_run("ove l'informativa sia fornita digitalmente "); r.bold = True; r.font.size = Pt(11)
r = p.add_run("(es. prima dell'accesso a un questionario su Microsoft "
              "Form, prima dell'iscrizione a qualche servizio, etc.) e:"); r.font.size = Pt(11)

p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.2)
r = p.add_run("o  "); r.font.size = Pt(11)
r = p.add_run("riguardi soggetti maggiorenni: "); r.bold = True; r.italic = True; r.font.size = Pt(11)
r = p.add_run(
    "sarà sufficiente una check box dal tenore \"Dichiaro di aver preso "
    "visione dell'informativa [nome informativa, inserendo il link "
    "all'informativa fornito dall'ufficio privacy al termine della "
    "redazione della stessa] e presto il consenso al trattamento ivi "
    "descritto\".")
r.font.size = Pt(11)

p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.2)
r = p.add_run("o  "); r.font.size = Pt(11)
r = p.add_run("riguardi soggetti minorenni in grado di comprendere l'informativa (sopra i 14 anni) "); r.bold = True; r.italic = True; r.font.size = Pt(11)
r = p.add_run("sarà necessario prevedere:"); r.font.size = Pt(11)

p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.8)
r = p.add_run(
    "▪  una check box destinata al minorenne, dal tenore \"Dichiaro di "
    "aver preso visione dell'informativa [nome informativa, inserendo "
    "il link all'informativa fornito dall'ufficio privacy al termine "
    "della redazione della stessa]\";")
r.font.size = Pt(11)

p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.8)
r = p.add_run(
    "▪  una check box destinata agli esercenti della potestà "
    "genitoriale, preceduta dalla richiesta di indicare nome e cognome "
    "(di entrambi) e un recapito (telefonico o e-mail), dal tenore "
    "\"Dichiariamo di aver preso visione dell'informativa [nome "
    "informativa, inserendo il link all'informativa fornito dall'ufficio "
    "privacy al termine della redazione della stessa] e prestiamo il "
    "consenso al trattamento ivi descritto\";")
r.font.size = Pt(11)

p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.2)
r = p.add_run("o  "); r.font.size = Pt(11)
r = p.add_run("riguardi soggetti minorenni: "); r.bold = True; r.italic = True; r.font.size = Pt(11)
r = p.add_run(
    "una check box destinata agli esercenti della potestà genitoriale, "
    "preceduta dalla richiesta di indicare nome e cognome (di entrambi) "
    "e un recapito telefonico, dal tenore \"Dichiariamo di aver preso "
    "visione dell'informativa [nome informativa, inserendo il link "
    "all'informativa fornito dall'ufficio privacy al termine della "
    "redazione della stessa] e prestiamo il consenso al trattamento ivi "
    "descritto\".")
r.font.size = Pt(11)

spacer(doc)
body(doc,
    "ATTENZIONE: tutte le check box dovranno essere bloccanti, ossia "
    "non dovrà essere possibile proseguire senza averle sbarrate.",
    italic=True)

spacer(doc)
p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
r = p.add_run("•  "); r.font.size = Pt(11)
r = p.add_run("ove l'informativa sia fornita cartacea, "); r.bold = True; r.font.size = Pt(11)
r = p.add_run(
    "il consenso andrà manifestato attraverso l'inserimento in calce "
    "all'informativa delle seguenti sezioni, a seconda delle diverse "
    "casistiche (soggetto maggiorenne, soggetto minorenne in grado di "
    "comprendere l'informativa e soggetto minorenne).")
r.font.size = Pt(11)

doc.add_page_break()

# ===================== SEZIONE B — ADULTS =====================
section(doc, "Sezione B")
spacer(doc)
hcenter(doc,
    "CONSENSO INFORMATO PER IL TRATTAMENTO DI DATI PERSONALI "
    "NELL'AMBITO DEI PROGETTI DI RICERCA",
    bold=True, size=11)
hcenter(doc, "soggetti maggiorenni", bold=True, italic=True, size=10)
spacer(doc)

body(doc, "Il/la sottoscritto/a ………………………………………………………………………  "
          "in relazione al Progetto ARISE")
spacer(doc)

checkbox(doc, "DICHIARA di aver preso visione dell'Informativa per il "
              "trattamento dei dati personali di cui alla sezione A")
spacer(doc)
checkbox(doc, "PRESTA IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i propri dati personali per le finalità e "
              "secondo le modalità ivi descritte")
spacer(doc)
checkbox(doc, "NEGA IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i propri dati personali per le finalità e "
              "secondo le modalità ivi descritte")
spacer(doc)
checkbox(doc, "AUTORIZZA l'Università degli Studi di Genova all'utilizzo "
              "delle proprie immagini e/o della propria voce per "
              "realizzazione di video e materiali multimediali, "
              "realizzati e utilizzati esclusivamente per finalità di "
              "ricerca scientifica")
spacer(doc)
checkbox(doc, "PRESTA IL CONSENSO affinché l'Università degli Studi di "
              "Genova utilizzi i dati di contatto per eventuali "
              "comunicazioni successive alla conclusione del progetto o "
              "affinché possa esprimere, se lo riterrà, un nuovo "
              "specifico consenso per una ulteriore ricerca")
spacer(doc)
checkbox(doc, "NEGA IL CONSENSO affinché l'Università degli Studi di "
              "Genova utilizzi i dati di contatto per eventuali "
              "comunicazioni successive alla conclusione del progetto o "
              "affinché possa esprimere, se lo riterrà, un nuovo "
              "specifico consenso per una ulteriore ricerca")

spacer(doc, 2)
body(doc, "Data ___/___/______")
spacer(doc)
body(doc, "Nome _________________________     Cognome _________________________")
spacer(doc)
body(doc, "Firma leggibile ________________________________________")

doc.add_page_break()

# ===================== SEZIONE B — MINORI =====================
hcenter(doc,
    "CONSENSO INFORMATO PER IL TRATTAMENTO DI DATI PERSONALI "
    "NELL'AMBITO DEI PROGETTI DI RICERCA",
    bold=True, size=11)
hcenter(doc, "da compilare nel caso di raccolta di dati personali di minori",
        bold=True, italic=True, size=10)
spacer(doc)

body(doc,
    "Nota: lo studio ARISE NON coinvolge soggetti minorenni. La presente "
    "sezione è inclusa per completezza del template e non è applicabile "
    "al presente studio.", italic=True)
spacer(doc)

body(doc,
    "I sottoscritti, titolari della potestà genitoriale di "
    "……………………………………… in relazione al Progetto ARISE")
spacer(doc)

checkbox(doc, "DICHIARANO di aver preso visione dell'Informativa per il "
              "trattamento dei dati personali di cui alla sezione A)")
spacer(doc)
checkbox(doc, "PRESTANO IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i dati personali del proprio figlio per le "
              "finalità e secondo le modalità ivi descritte")
spacer(doc)
checkbox(doc, "NEGANO IL CONSENSO affinché l'Università degli Studi di "
              "Genova tratti i dati personali del proprio figlio per le "
              "finalità e secondo le modalità ivi descritte")
spacer(doc)
checkbox(doc, "AUTORIZZANO l'Università degli Studi di Genova all'utilizzo "
              "delle immagini e/o della voce del proprio figlio per la "
              "realizzazione di video e materiali multimediali, "
              "realizzati e utilizzati esclusivamente per finalità di "
              "ricerca scientifica")
spacer(doc)
checkbox(doc, "PRESTANO IL CONSENSO affinché l'Università degli Studi di "
              "Genova utilizzi i dati di contatto per eventuali "
              "comunicazioni successive alla conclusione del progetto")
spacer(doc)
checkbox(doc, "NEGANO IL CONSENSO affinché l'Università degli Studi di "
              "Genova utilizzi i dati di contatto per eventuali "
              "comunicazioni successive alla conclusione del progetto")

spacer(doc, 2)
body(doc, "Data ___/___/______")
spacer(doc)
body(doc, "Nome ____________________   Cognome ____________________")
spacer(doc)
body(doc, "Firma leggibile ________________________________________")
spacer(doc)
body(doc, "Nome ____________________   Cognome ____________________")
spacer(doc)
body(doc, "Firma leggibile ________________________________________")

doc.save(OUT)
print(f"Wrote {OUT}")
