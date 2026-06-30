"""CERA Nota Informativa + Consenso for the ARISE study, in plain Q&A
style mirroring the prior DINOGMI/Testa submission. Single UNIGE site
(Laboratori del Dipartimento DINOGMI, Campus di Savona). Only
Nota Informativa + Consenso participation;
the GDPR Informativa Privacy lives in a separate document
(tools_build_cera_privacy.py).
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2")
OUT  = ROOT / "ARISE_CERA_Nota_Informativa_Consenso.docx"

STUDY_TITLE = (
    "ARISE — A feasibility study on markerless video-based biomechanical "
    "assessment of the Sit-to-Stand task in healthy adults and older adults"
)

def hcenter(doc, t, *, bold=False, size=11):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)

def section(doc, t):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = True; r.font.size = Pt(13)

def label(doc, t, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(11)

def body(doc, t, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.italic = italic; r.font.size = Pt(11)

def qa(doc, q, a):
    p = doc.add_paragraph()
    rq = p.add_run(q + " "); rq.bold = True; rq.italic = True; rq.font.size = Pt(11)
    ra = p.add_run(a); ra.font.size = Pt(11)

def bullet(doc, t, indent=0.6):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run("•  " + t); r.font.size = Pt(11)

def spacer(doc, n=1):
    for _ in range(n): doc.add_paragraph()

doc = Document()
s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)

# ===================== HEADER =====================
hcenter(doc, "UNIVERSITÀ DEGLI STUDI DI GENOVA", bold=True, size=14)
hcenter(doc, "Dipartimento di Neuroscienze, Riabilitazione, Oftalmologia, "
             "Genetica e Scienze Materno-Infantili")
spacer(doc)

# ===================== SEZIONE A — NOTA INFORMATIVA =====================
section(doc, "Sezione A – Nota informativa sullo studio")
spacer(doc)

body(doc, "Gentile Signora/Egregio Signore,")
body(doc, f"Le è stato chiesto di partecipare ad uno studio dal titolo \"{STUDY_TITLE}\".")
body(doc,
    "Prima che Lei prenda una decisione in merito, è importante che "
    "comprenda il motivo dello studio e cosa Le sarà chiesto di fare, "
    "qualora decidesse di prendervi parte.")
body(doc,
    "Lo sperimentatore ed i suoi collaboratori, oltre alle spiegazioni "
    "che Le forniranno durante questo colloquio, sono a Sua completa "
    "disposizione per qualsiasi chiarimento.")
body(doc,
    "Questo documento ha lo scopo di fornirle un'informazione corretta "
    "e completa affinché Lei possa esprimere una scelta libera e "
    "consapevole. Inoltre, qualora lo desiderasse, prima di decidere, "
    "può chiedere un parere a qualunque persona di sua fiducia.")
body(doc,
    "Il responsabile dello studio è Marco Testa, Professore Associato "
    "presso il Dipartimento di Neuroscienze, Riabilitazione, "
    "Oftalmologia, Genetica e Scienze Materno-Infantili (DINOGMI).")

spacer(doc)
hcenter(doc, "NOTA INFORMATIVA", bold=True, size=12)
spacer(doc)

qa(doc, "Qual è lo scopo dello studio?",
   "Lo studio ha l'obiettivo di (1) raccogliere dati biomeccanici sul "
   "Sit-to-Stand (passaggio dalla posizione seduta a quella eretta) "
   "in una popolazione di adulti (>18 anni); (2) caratterizzare i "
   "principali indicatori cinematici del movimento (durata di fase, "
   "velocità del bacino, angoli articolari di tronco e ginocchio, "
   "simmetria laterale); (3) validare i dati ottenuti mediante "
   "telecamere RGB e/o RGB-D senza marker confrontandoli con misure "
   "di riferimento ottenute mediante sistemi di analisi cinematica e "
   "cinetica del movimento gold-standard validati con metodologie "
   "non invasive.")

qa(doc, "È obbligato a partecipare?",
   "La partecipazione è totalmente volontaria ed il Suo contributo "
   "aiuterà a raggiungere gli scopi dello studio.")

qa(doc, "Cosa accadrà se decide di partecipare allo studio?",
   "Accettando di partecipare Le chiederemo preliminarmente di "
   "fornire alcune informazioni sociodemografiche (età, altezza, "
   "peso, sesso, genere) e anamnestiche generiche sullo stato di "
   "salute, oltre ai Suoi dati di contatto (nome, cognome, recapito "
   "telefonico, e-mail). Questi ultimi saranno utilizzati "
   "esclusivamente per consentirLe di esercitare i Suoi diritti "
   "GDPR (richiesta di cancellazione, restituzione dei risultati "
   "personali) e saranno custoditi unicamente dal Coordinatore del "
   "progetto. Successivamente parteciperà alla sperimentazione "
   "presso i Laboratori di ricerca del Dipartimento DINOGMI, ubicati "
   "al Campus Universitario di Savona.")

qa(doc, "Quale sarà il suo impegno? Cosa dovrà fare?",
   "Dovrà eseguire alcune ripetizioni del Sit-to-Stand su una sedia "
   "standard, con le braccia incrociate al petto, secondo i "
   "protocolli standardizzati 5xSTS (cinque ripetizioni a tempo) e "
   "30-second Chair Stand Test, sempre alla presenza di un operatore "
   "qualificato del gruppo di ricerca. Prima di iniziare verrà "
   "eseguita una breve fase di calibrazione (circa 5 secondi). "
   "L'esecuzione del movimento sarà registrata mediante telecamere "
   "RGB e/o RGB-D posizionate davanti a Lei. In una parte facoltativa "
   "dello studio, la registrazione video potrà essere effettuata "
   "contemporaneamente a sistemi di analisi cinematica e cinetica del "
   "movimento gold-standard validati con metodologie non invasive.")

qa(doc, "Quali potrebbero essere i rischi, gli effetti collaterali, i disagi?",
   "I rischi sono minimi e sovrapponibili a quelli di una normale "
   "attività motoria leggera: possibile lieve affaticamento muscolare "
   "dopo le ripetizioni, lieve indolenzimento muscolare nelle ore "
   "successive (fisiologico), e basso rischio residuo di caduta, "
   "mitigato dalla supervisione continua dell'operatore accanto a Lei "
   "durante ogni ripetizione e dai criteri di inclusione che "
   "escludono soggetti con instabilità o controindicazioni mediche. "
   "Non sono previsti rischi di tipo psicologico, farmacologico, "
   "chirurgico né invasivo.")

qa(doc, "Potrà cambiare idea dopo aver accettato di partecipare?",
   "Certo, il consenso a partecipare potrà essere ritirato in qualsiasi "
   "momento e senza dover fornire alcuna spiegazione. In questo caso i "
   "Suoi dati non saranno salvati.")

qa(doc, "Quanto dura l'impegno del partecipante alla ricerca?",
   "La sperimentazione in laboratorio non durerà più di 10-15 minuti "
   "per sessione. Le verrà chiesto di partecipare a 2 sessioni "
   "complessive nell'arco di una settimana. L'eventuale sessione di "
   "confronto con la strumentazione di riferimento è facoltativa ed "
   "è ricompresa entro la stessa durata.")

qa(doc, "Cosa succede se decide di non partecipare allo studio?", "Nulla.")

qa(doc, "Informazioni circa i risultati dello studio:",
   "I risultati saranno resi noti in forma anonima ed aggregata dopo la "
   "conclusione dello studio. Potranno essere ricercati in modo "
   "autonomo da chiunque fosse interessato dopo la loro pubblicazione "
   "oppure richiesti agli sperimentatori.")

spacer(doc)
body(doc,
    "Per ulteriori informazioni e comunicazioni durante lo studio sarà "
    "a disposizione il seguente personale:")
contacts = [
    ("Prof. Marco Testa", "+39 331 2611548", "marco.testa@unige.it"),
    ("Dott. Mirko Job",   "+39 349 5199778", "mirko.job.1991@gmail.com"),
    ("Paolo Pavani (Innovina S.r.l.)",   "[telefono da inserire]", "[email da inserire]"),
    ("Hadil Sahraoui (Innovina S.r.l.)", "[telefono da inserire]", "[email da inserire]"),
]
for name, phone, email in contacts:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f"{name}   ·   Cellulare: {phone}   ·   {email}")
    r.font.size = Pt(11)

doc.add_page_break()

# ===================== SEZIONE B — CONSENSO PARTECIPAZIONE =====================
hcenter(doc, "UNIVERSITÀ DEGLI STUDI DI GENOVA", bold=True, size=14)
hcenter(doc, "Dipartimento di Neuroscienze, Riabilitazione, Oftalmologia, "
             "Genetica e Scienze Materno-Infantili")
spacer(doc)

section(doc, "Sezione B – Consenso alla partecipazione allo studio "
             "(soggetti maggiorenni)")
spacer(doc)

label(doc, "1) Confermo di:")
bullet(doc, f"aver ricevuto spiegazioni esaustive in merito allo studio "
            f"dal titolo \"{STUDY_TITLE}\"")
bullet(doc, "aver preso visione della nota informativa relativa allo "
            "studio di cui alla Sezione A e di averne ricevuto copia")
bullet(doc, "aver avuto l'opportunità di fare domande in merito allo studio.")

spacer(doc)
label(doc, "2) Sono consapevole:")
bullet(doc, "dei rischi e dei benefici che possono derivare dalla "
            "partecipazione a questo studio")
bullet(doc, "che la mia partecipazione è volontaria, e di essere libero "
            "di potermi ritirare in qualunque momento senza dover dar "
            "spiegazioni e senza che la mia assistenza medica o i miei "
            "diritti ne siano condizionati.")

spacer(doc)
label(doc, "3) Accetto di partecipare a questo studio")

spacer(doc, 2)
body(doc, "Nome _________________________     Cognome _________________________")
spacer(doc)
body(doc, "Firma ________________________________________     Data ___/___/______")

spacer(doc, 2)
body(doc, "Nome dello sperimentatore che ha raccolto il consenso "
          "________________________________________")
spacer(doc)
body(doc, "Firma dello sperimentatore che ha raccolto il consenso "
          "________________________________________")

doc.save(OUT)
print(f"Wrote {OUT}")
