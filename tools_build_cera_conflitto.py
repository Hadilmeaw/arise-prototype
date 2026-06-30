"""CERA Dichiarazione sul Conflitto di Interessi for the ARISE study,
matching the prior Testa/DINOGMI submissions: concise, signed by
Prof. Marco Testa, defaulted to "no conflict" with the structural
academic-industry collaboration with Innovina S.r.l. disclosed for
full transparency.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2")
OUT  = ROOT / "ARISE_CERA_Dichiarazione_Conflitto_Interessi.docx"

CHECKED, UNCHECKED = "☒", "☐"

def hcenter(doc, t, *, bold=False, size=11):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)

def hright(doc, t, *, bold=False, size=11):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)

def body(doc, t, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.italic = italic; r.font.size = Pt(11)

def spacer(doc, n=1):
    for _ in range(n): doc.add_paragraph()

doc = Document()
s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)

hcenter(doc, "UNIVERSITÀ DEGLI STUDI DI GENOVA", bold=True, size=14)
hcenter(doc, "DIPARTIMENTO DI NEUROSCIENZE, RIABILITAZIONE, OFTALMOLOGIA, "
             "GENETICA E SCIENZE MATERNO-INFANTILI (DINOGMI)")
spacer(doc)

hright(doc,
    "Al Presidente\nComitato Etico per la Ricerca di Ateneo\n"
    "Università degli Studi di Genova")
spacer(doc)

hcenter(doc, "DICHIARAZIONE SUL CONFLITTO DI INTERESSI",
        bold=True, size=14)
spacer(doc, 1)

body(doc,
    "In relazione allo svolgimento del progetto "
    "\"ARISE — A feasibility study on markerless video-based "
    "biomechanical assessment of the Sit-to-Stand task in healthy adults "
    "and older adults\" (acronimo: ARISE),")

body(doc,
    "finanziato dal progetto ARISE nell'ambito del Fondo Europeo di "
    "Sviluppo Regionale (FESR/ERDF), programma HealthTech regionale "
    "della Regione Liguria, con Innovina S.r.l. quale partner tecnico-"
    "industriale del progetto, per cui chiede l'espressione del parere "
    "da parte del CERA,")

body(doc,
    "il sottoscritto Prof. Marco Testa, Professore Associato presso il "
    "Dipartimento di Neuroscienze, Riabilitazione, Oftalmologia, "
    "Genetica e Scienze Materno-Infantili (DINOGMI) dell'Università "
    "degli Studi di Genova, responsabile del progetto, dichiara di:")

spacer(doc)

p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
r = p.add_run(f"{CHECKED}  non avere conflitti di interesse"); r.font.size = Pt(11)

p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
r = p.add_run(f"{UNCHECKED}  di avere conflitti di interesse"); r.font.size = Pt(11)

spacer(doc)
body(doc, "Specificare ……………………………………………………………………………………………………")

spacer(doc, 2)

body(doc,
    "Per piena trasparenza, si dà atto che il progetto ARISE è "
    "realizzato nell'ambito di una collaborazione di ricerca tra "
    "l'Università degli Studi di Genova (DINOGMI e laboratorio REHELAB) "
    "e Innovina S.r.l. quale partner tecnico-industriale che fornisce "
    "il software di acquisizione e analisi video utilizzato nello "
    "studio. Tale collaborazione è regolata da apposito accordo "
    "nell'ambito del programma di finanziamento FESR/ERDF Liguria "
    "HealthTech ed è già dichiarata nella Sezione 1 (Componenti del "
    "gruppo di ricerca) e nella Sezione 2 (Eventuali enti finanziatori "
    "esterni o Sponsor) della Richiesta di Parere CERA. Il sottoscritto "
    "non riceve compensi personali, partecipazioni azionarie, onorari "
    "di consulenza o altri benefici economici diretti da Innovina "
    "S.r.l. al di fuori dell'attività istituzionale di ricerca regolata "
    "dall'accordo summenzionato.",
    italic=True)

spacer(doc, 3)
body(doc, "Genova, ___/___/______")

spacer(doc, 4)
hcenter(doc, "Prof. Marco Testa PT, PhD", bold=True)
hcenter(doc, "Delegato del Rettore per il Campus Universitario di Savona")
hcenter(doc, "Presidente del Master in Riabilitazione dei Disordini "
             "Muscoloscheletrici")
hcenter(doc, "Università di Genova, Campus of Savona")

doc.save(OUT)
print(f"Wrote {OUT}")
