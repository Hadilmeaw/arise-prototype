"""ARISE master document template.

Single, reusable template applied to every project deliverable so the
documentation set has a consistent look.

Template features:
  - Cover page with accent strip, project mark, document title and metadata
  - Page header (top of every body page): document title left, page number right
  - Page footer: organisation name centred
  - Calibri 11pt body, navy headings, plain black-border tables with navy
    header rows

Public API:
  build_arise_document(
      input_md, output_docx, doc_title, doc_subtitle=None,
      meta=None, doc_short=None
  )
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Palette ────────────────────────────────────────────────────────

NAVY = RGBColor(0x1A, 0x33, 0x5C)
INK  = RGBColor(0x14, 0x14, 0x14)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
PALE = RGBColor(0xD9, 0xD9, 0xD9)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ─── Markdown parsing helpers ───────────────────────────────────────

LINK_RE = re.compile(r"\[(.+?)\]\((.+?)\)")


def _add_inline_runs(p, text):
    pattern = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\)|\*[^*]+?\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = p.add_run(token[2:-2]); r.bold = True
        elif token.startswith("`"):
            r = p.add_run(token[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(10)
        elif token.startswith("["):
            lm = LINK_RE.match(token)
            r = p.add_run(lm.group(1))
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1); r.underline = True
        elif token.startswith("*"):
            r = p.add_run(token[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def _set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _add_md_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    has_header = any(c.strip() for c in rows[0])
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    if not has_header:
        rows = rows[1:]
        tbl._element.remove(tbl.rows[0]._tr)
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text.strip())
            if has_header and i == 0:
                for run in p.runs:
                    run.bold = True
                _set_cell_shading(cell, "1A335C")
                for run in p.runs:
                    run.font.color.rgb = WHITE
            elif not has_header and j == 0:
                for run in p.runs:
                    run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def _parse_table_block(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line):
            i += 1; continue
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "": cells = cells[1:]
        if cells and cells[-1] == "": cells = cells[:-1]
        rows.append(cells)
        i += 1
    return rows, i


IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.+?)\)$")


def _render_markdown(doc, md_text, *, skip_first_h1=True):
    """Render markdown into the doc. The first H1 (which is normally the
    document title) is skipped because the cover page already carries
    the title."""
    from docx.shared import Inches as _Inches
    lines = md_text.splitlines()
    i, n = 0, len(lines)
    first_h1_seen = not skip_first_h1
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1; continue
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            i += 1; continue
        # Image (one image per line)
        im = IMAGE_RE.match(stripped)
        if im:
            caption, path = im.group(1), im.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = p.add_run()
                run.add_picture(path, width=_Inches(6.0))
            except Exception as e:
                run = p.add_run(f"[Image not found: {path}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0xC0, 0x2A, 0x2A)
            if caption:
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_after = Pt(10)
                cr = cap_p.add_run(caption)
                cr.font.name = "Calibri"
                cr.font.size = Pt(9)
                cr.font.italic = True
                cr.font.color.rgb = GREY
            i += 1; continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            ttl = m.group(2).strip()
            ttl = re.sub(r"^\*\*(.+)\*\*$", r"\1", ttl)
            if level == 1 and not first_h1_seen:
                first_h1_seen = True
                i += 1; continue
            p = doc.add_heading(level=min(level, 4))
            _add_inline_runs(p, ttl)
            i += 1; continue
        if stripped.startswith("|"):
            rows, i = _parse_table_block(lines, i)
            _add_md_table(doc, rows)
            continue
        if re.match(r"^[-*+]\s+", stripped):
            while i < n and re.match(r"^[-*+]\s+", lines[i].strip()):
                item = re.sub(r"^[-*+]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, item)
                i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, item)
                i += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            r = p.add_run(stripped[2:])
            r.italic = True; r.font.color.rgb = GREY
            i += 1; continue
        if stripped.startswith("```"):
            i += 1; code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code))
            r.font.name = "Consolas"; r.font.size = Pt(9)
            continue
        para_lines = [line]; i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#|\||[-*+]\s|\d+\.\s|>\s|```|-{3,})", lines[i].strip()
        ):
            para_lines.append(lines[i]); i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        _add_inline_runs(p, para_text)


# ─── Template construction ──────────────────────────────────────────

def _apply_global_styles(doc):
    """Set up document-wide styles."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    # Body text
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Headings
    for lvl, size in [(1, 18), (2, 14), (3, 12), (4, 11)]:
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = NAVY
        s.font.bold = True


def _set_header_footer(section, doc_short):
    """Page header (doc title + page number) and centred footer."""
    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Left text
    r1 = hp.add_run(f"ARISE   |   {doc_short}")
    r1.font.name = "Calibri"; r1.font.size = Pt(9); r1.font.color.rgb = GREY
    # Tab to right + page-number field
    r2 = hp.add_run("\t\t")
    r3 = hp.add_run("Page ")
    r3.font.name = "Calibri"; r3.font.size = Pt(9); r3.font.color.rgb = GREY
    page_run = hp.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    page_run._r.append(fld1); page_run._r.append(instr); page_run._r.append(fld2)
    page_run.font.name = "Calibri"; page_run.font.size = Pt(9); page_run.font.color.rgb = GREY

    # Subtle bottom border on header
    pPr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'D9D9D9')
    pbdr.append(bottom)
    pPr.append(pbdr)

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Innovina S.r.l.")
    r.font.name = "Calibri"; r.font.size = Pt(9); r.font.color.rgb = GREY


def _build_cover_page(doc, doc_title, doc_subtitle, meta):
    """Cover page. Left accent strip, project mark, document title,
    optional subtitle, metadata table."""
    # Accent bar (top-left)
    # Implemented as an empty paragraph followed by visual elements

    # Top spacer
    for _ in range(2):
        doc.add_paragraph()

    # Project mark
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ARISE")
    r.font.name = "Calibri"; r.font.size = Pt(40); r.font.bold = True
    r.font.color.rgb = NAVY

    # Tagline
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(40)
    r2 = p2.add_run("Augmented Rehabilitation & Intelligent System for Enhancement")
    r2.font.name = "Calibri"; r2.font.size = Pt(12); r2.font.italic = True
    r2.font.color.rgb = GREY

    # Document title
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(4)
    r3 = p3.add_run(doc_title)
    r3.font.name = "Calibri"; r3.font.size = Pt(26); r3.font.bold = True
    r3.font.color.rgb = INK

    if doc_subtitle:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(40)
        r4 = p4.add_run(doc_subtitle)
        r4.font.name = "Calibri"; r4.font.size = Pt(13); r4.font.italic = True
        r4.font.color.rgb = GREY
    else:
        doc.add_paragraph()
        doc.add_paragraph()

    # Metadata table
    if meta:
        tbl = doc.add_table(rows=len(meta), cols=2)
        tbl.style = "Light List Accent 1"
        for i, (k, v) in enumerate(meta.items()):
            c1 = tbl.rows[i].cells[0]; c1.text = ""
            pp = c1.paragraphs[0]
            rr = pp.add_run(k)
            rr.font.name = "Calibri"; rr.font.size = Pt(10); rr.font.bold = True
            rr.font.color.rgb = INK
            c2 = tbl.rows[i].cells[1]; c2.text = ""
            pp2 = c2.paragraphs[0]
            rr2 = pp2.add_run(v)
            rr2.font.name = "Calibri"; rr2.font.size = Pt(10)
            rr2.font.color.rgb = INK


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ─── Public API ─────────────────────────────────────────────────────

def build_arise_document(
    input_md: Path,
    output_docx: Path,
    *,
    doc_title: str,
    doc_subtitle: str = None,
    meta: dict = None,
    doc_short: str = None,
):
    """Build an ARISE-styled docx from a markdown content file."""
    doc = Document()
    _apply_global_styles(doc)

    section = doc.sections[0]
    _set_header_footer(section, doc_short or doc_title)

    _build_cover_page(doc, doc_title, doc_subtitle, meta)
    _page_break(doc)

    md_text = Path(input_md).read_text(encoding="utf-8")
    _render_markdown(doc, md_text)

    doc.save(str(output_docx))


if __name__ == "__main__":
    # Allow CLI use for quick rebuilds:
    #   python tools_arise_doc_template.py <input.md> <output.docx> "<title>" "<subtitle>"
    import sys
    if len(sys.argv) < 4:
        print("Usage: tools_arise_doc_template.py input.md output.docx 'title' ['subtitle']")
        sys.exit(2)
    inp, out, ttl = sys.argv[1], sys.argv[2], sys.argv[3]
    sub = sys.argv[4] if len(sys.argv) > 4 else None
    build_arise_document(Path(inp), Path(out), doc_title=ttl, doc_subtitle=sub)
    print(f"Wrote {out}")
