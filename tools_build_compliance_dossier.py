"""Build the combined ARISE Compliance Dossier as a single .docx.

The dossier contains three documents in sequence:
  1. Data Management Plan
  2. Data Protection Impact Assessment
  3. MDR Compliance Plan

Each document keeps its own internal structure. The combined file has a
master cover, internal navigation table, page breaks between documents,
and consistent header/footer throughout.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Palette
NAVY = RGBColor(0x1A, 0x33, 0x5C)
INK  = RGBColor(0x14, 0x14, 0x14)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
LIGHT = RGBColor(0xF2, 0xF2, 0xF2)
PALE = RGBColor(0xD9, 0xD9, 0xD9)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ----- markdown parsing helpers (reused from tools_md_to_docx.py) -----

LINK_RE = re.compile(r"\[(.+?)\]\((.+?)\)")


def add_runs(p, text):
    pattern = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = p.add_run(token[2:-2]); r.bold = True
        elif token.startswith("`"):
            r = p.add_run(token[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        elif token.startswith("["):
            lm = LINK_RE.match(token)
            r = p.add_run(lm.group(1))
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            r.underline = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_md_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    has_header = any(c.strip() for c in rows[0])
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    if not has_header:
        rows = rows[1:]
        tbl._element.remove(tbl.rows[0]._tr)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, val.strip())
            if has_header and i == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "1A335C")
                for run in p.runs:
                    run.font.color.rgb = WHITE
            elif not has_header and j == 0:
                for run in p.runs:
                    run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def parse_table_block(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line):
            i += 1; continue
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "": cells = cells[1:]
        if cells and cells[-1] == "": cells = cells[:-1]
        rows.append(cells)
        i += 1
    return rows, i


def render_markdown(doc, md_text):
    """Render markdown text into the doc, skipping the H1 title and any
    front-matter cover-block table (the combined document has its own
    cover and per-document hero pages)."""
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    skipped_title = False
    skipped_front_matter_table = False

    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1; continue
        # Skip horizontal rules
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            i += 1; continue
        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            ttl = m.group(2).strip()
            ttl = re.sub(r"^\*\*(.+)\*\*$", r"\1", ttl)
            if level == 1 and not skipped_title:
                # Skip the first H1 (document title; we render it on the hero page)
                skipped_title = True
                i += 1; continue
            if level == 2 and ttl.lower().startswith("how this document"):
                # Skip the "How this document is organised" subsection;
                # the master TOC carries that
                # actually we keep it, it's useful
                pass
            p = doc.add_heading(level=min(level, 4))
            # Re-apply our heading style colors (set globally below)
            add_runs(p, ttl)
            i += 1; continue
        # Tables
        if stripped.startswith("|"):
            # The first table after H1 in a doc is typically the cover-block
            # metadata table. Skip the very first one we encounter if we haven't yet.
            rows, j = parse_table_block(lines, i)
            if not skipped_front_matter_table and len(rows) >= 3 \
               and rows[0] and rows[0][0].strip() == "" and len(rows[0]) == 2:
                # empty-header two-column => cover-block
                skipped_front_matter_table = True
                i = j; continue
            add_md_table(doc, rows)
            i = j; continue
        # Bullet list
        if re.match(r"^[-*+]\s+", stripped):
            while i < n and re.match(r"^[-*+]\s+", lines[i].strip()):
                item = re.sub(r"^[-*+]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, item)
                i += 1
            continue
        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_runs(p, item)
                i += 1
            continue
        # Blockquote
        if stripped.startswith("> "):
            content = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            r = p.add_run(content)
            r.italic = True
            r.font.color.rgb = GREY
            i += 1; continue
        # Code block
        if stripped.startswith("```"):
            i += 1; code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code))
            r.font.name = "Consolas"; r.font.size = Pt(9)
            continue
        # Paragraph
        para_lines = [line]; i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#|\||[-*+]\s|\d+\.\s|>\s|```|-{3,})", lines[i].strip()
        ):
            para_lines.append(lines[i]); i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        add_runs(p, para_text)


# ----- chrome: cover, footer, page break, hero pages -----

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_section_break(doc):
    """Insert a section break before next page so header/footer can change."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    sectType = OxmlElement('w:type'); sectType.set(qn('w:val'), 'nextPage')
    sectPr.append(sectType)
    pPr.append(sectPr)


def set_page_footer(section, text_left, page_label):
    footer = section.footer
    # Clear any existing paragraphs
    for p in list(footer.paragraphs):
        if p.text:
            for r in p.runs: r.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Left side: doc identifier
    r1 = fp.add_run(text_left)
    r1.font.size = Pt(9); r1.font.color.rgb = GREY
    # Tab to center, then page label
    r2 = fp.add_run("\t\t")
    r3 = fp.add_run(page_label + "  ")
    r3.font.size = Pt(9); r3.font.color.rgb = GREY
    # Page number field
    run = fp.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.size = Pt(9); run.font.color.rgb = GREY


def add_text(doc, text, *, size=11, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color
    return p


def build_cover_page(doc):
    """Master cover page for the dossier."""
    # Big spacing top
    for _ in range(3):
        doc.add_paragraph()
    add_text(doc, "ARISE", size=48, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_text(doc,
             "Augmented Rehabilitation & Intelligent System for Enhancement",
             size=14, color=GREY, italic=True, space_after=20)
    add_text(doc, "Compliance Dossier", size=32, bold=True, color=INK,
             space_after=4)
    add_text(doc, "Phase 1, TRL6 demonstration. Class IIa under MDR Annex VIII Rule 11", size=14, color=GREY,
             italic=True, space_after=40)

    # Three documents enumerated
    add_text(doc, "Documents in this dossier", size=12, bold=True, color=NAVY,
             space_after=8)
    docs = [
        ("01", "Data Management Plan (DMP)",
         "Engineering specification: data architecture, lifecycle, schemas, pipelines, quality, access"),
        ("02", "Data Protection Impact Assessment (DPIA)",
         "GDPR Article 35 risk assessment, lawful bases, mitigations, twelve-risk register"),
        ("03", "MDR Compliance Plan",
         "Confirmed Class IIa under Rule 11. Annex IX conformity-assessment route. Clinical investigation under MDR Articles 62-82 in Phase 1, Notified Body engagement in Phase 2"),
    ]
    for num, title, sub in docs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r1 = p.add_run(num + "   ")
        r1.font.name = "Calibri"; r1.font.size = Pt(14); r1.font.bold = True
        r1.font.color.rgb = NAVY
        r2 = p.add_run(title)
        r2.font.name = "Calibri"; r2.font.size = Pt(14); r2.font.bold = True
        r2.font.color.rgb = INK
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.9)
        p2.paragraph_format.space_after = Pt(10)
        r3 = p2.add_run(sub)
        r3.font.name = "Calibri"; r3.font.size = Pt(11); r3.font.color.rgb = GREY

    # Footer-area dossier metadata
    for _ in range(2):
        doc.add_paragraph()
    add_text(doc, "Project context", size=10, bold=True, color=NAVY)
    meta = [
        ("Funding instrument", "ERDF"),
        ("Project duration", "18 months, June 2026 to December 2027"),
        ("Work Package", "WP1, Task T1.4 (Compliance Management)"),
        ("Prepared by", "Innovina S.r.l."),
        ("Scientific partner", "DINOGMI, University of Genoa"),
        ("Clinical partner", "Studio Buccarella"),
        ("Confirmed MDR class", "Class IIa under Annex VIII Rule 11"),
        ("Conformity assessment route", "Annex IX (Notified Body required)"),
        ("Dossier version", "2.0, 9 June 2026"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Light List Accent 1"
    for i, (k, v) in enumerate(meta):
        c1 = tbl.rows[i].cells[0]
        c1.text = ""
        p = c1.paragraphs[0]
        r = p.add_run(k)
        r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = INK
        c2 = tbl.rows[i].cells[1]
        c2.text = ""
        p = c2.paragraphs[0]
        r = p.add_run(v)
        r.font.name = "Calibri"; r.font.size = Pt(10)
        r.font.color.rgb = INK


def build_master_toc(doc):
    add_page_break(doc)
    add_text(doc, "Contents", size=24, bold=True, color=NAVY, space_after=20)

    items = [
        ("01", "Data Management Plan"),
        ("02", "Data Protection Impact Assessment"),
        ("03", "MDR Compliance Plan"),
    ]
    for num, title in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r1 = p.add_run(num + "    ")
        r1.font.name = "Calibri"; r1.font.size = Pt(14); r1.font.bold = True
        r1.font.color.rgb = NAVY
        r2 = p.add_run(title)
        r2.font.name = "Calibri"; r2.font.size = Pt(14); r2.font.color.rgb = INK

    # Note about cross-references
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    r = p.add_run(
        "The three documents in this dossier cross-reference each other. "
        "Section numbers refer to the document within which the reference appears. "
        "Where one document defers detail to another, the referenced document is named explicitly."
    )
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.italic = True
    r.font.color.rgb = GREY


def build_hero_page(doc, number, title, subtitle, summary_lines):
    """Open a new document section: page break, then a centered hero page."""
    add_page_break(doc)
    for _ in range(3):
        doc.add_paragraph()
    # Big document number on the left
    add_text(doc, number, size=72, bold=True, color=NAVY, space_after=4)
    add_text(doc, title, size=28, bold=True, color=INK, space_after=4)
    add_text(doc, subtitle, size=14, color=GREY, italic=True, space_after=24)
    add_text(doc, "Document summary", size=11, bold=True, color=NAVY,
             space_after=6)
    for ln in summary_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("•  " + ln)
        r.font.name = "Calibri"; r.font.size = Pt(11); r.font.color.rgb = INK
    add_page_break(doc)


# ----- main build -----

def build(out_path: Path):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for lvl, size in [(1, 18), (2, 14), (3, 12), (4, 11)]:
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = NAVY
        s.font.bold = True

    # Footer with page numbers
    set_page_footer(section, "ARISE Compliance Dossier", "Page")

    # ===== Master cover =====
    build_cover_page(doc)

    # ===== Master TOC =====
    build_master_toc(doc)

    # ===== Document 01: DMP =====
    build_hero_page(
        doc,
        number="01",
        title="Data Management Plan",
        subtitle="Engineering specification for ARISE data",
        summary_lines=[
            "Four storage tiers: edge, hot, warm, cold.",
            "Three lifecycle blocks: Phase A (M6 to M15, WP3), Bridge (M8 to M18, WP4), Phase B (M19 to M24, WP5).",
            "Five top-level datasets: D-CORPUS (training corpus), D-TRIAL (trial operational), D-USABILITY, D-TELEMETRY, D-ENG.",
            "Schemas, pipelines split by phase, quality gates, access control roles.",
            "Compliance traceability matrix mapping engineering controls to GDPR and MDR.",
        ],
    )
    dmp_md = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2\ARISE_Data_Management_Plan.md").read_text(encoding="utf-8")
    render_markdown(doc, dmp_md)

    # ===== Document 02: DPIA =====
    build_hero_page(
        doc,
        number="02",
        title="Data Protection Impact Assessment",
        subtitle="GDPR Article 35 mandatory assessment",
        summary_lines=[
            "DPIA mandate under GDPR Article 35(3) trigger (a) automated decision-making, (b) special category at scale, (c) novel technology.",
            "Lawful bases summary across Article 6 and Article 9, with explicit non-reliance on 6(1)(d to f) and 9(2)(g, i).",
            "Twelve identified risks scored on likelihood and severity, mapped to project phases.",
            "Inline operational controls: retention schedule, deletion procedure, audit logging, breach response, encryption.",
            "Article 80 of MDR applies in Phase 1 (clinical investigation), Articles 83 to 87 take over in Phase 2 for the confirmed Class IIa device.",
        ],
    )
    dpia_md = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2\ARISE_DPIA.md").read_text(encoding="utf-8")
    render_markdown(doc, dpia_md)

    # ===== Document 03: MDR =====
    build_hero_page(
        doc,
        number="03",
        title="MDR Compliance Plan",
        subtitle="Medical-device regulatory alignment under EU 2017/745. Confirmed Class IIa",
        summary_lines=[
            "Classification confirmed as Class IIa under MDR Annex VIII Rule 11. The companion document ARISE Classification Scenarios is now archived.",
            "Conformity-assessment route: Annex IX (full QMS assessment plus assessment of the technical documentation). Notified Body involvement required.",
            "Phase 1 clinical investigation under MDR Articles 62 to 82, AE recording per Article 80.",
            "Phase 2 post-market obligations include Article 86 PSUR every two years, Article 87 vigilance, mandatory PMCF per Annex XIV Part B.",
            "Full ISO 13485 quality management system and ISO 14971 risk management process scoped for Class IIa.",
            "Notified Body engagement plan: shortlist at M14, pre-submission at M18 to M20, contract Q1 Phase 2, QMS audit Q2 Phase 2.",
        ],
    )
    mdr_md = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2\ARISE_MDR_Compliance_Plan.md").read_text(encoding="utf-8")
    render_markdown(doc, mdr_md)

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    out = Path(r"\\wsl.localhost\Ubuntu-20.04\home\ev04\v2\ARISE_Compliance_Dossier.docx")
    build(out)
